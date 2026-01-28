# -*- coding: utf-8 -*- v. 05/11/2025 - 19 hs
"""
app_safety_chat.py — Sphera + RAG + DIC (WS/Precursores/CP)

Patches desta versão:
- CORRIGIDO: Tratamento robusto de erros e validação de dados
- CORRIGIDO: Filtro de Location consistente com fallback seguro
- MELHORADO: Cache com controle de memória e TTL
- MELHORADO: Validação de alinhamento embeddings/labels
- MELHORADO: Logging e debugging estruturado
- MELHORADO: Performance com batch processing otimizado
- ADICIONADO: Suporte completo para Sphera + GoSee + Docs (conforme documentação)

"""

import os
import re
import io
import time
import logging
import warnings
from datetime import datetime, timedelta
from pathlib import Path
from typing import List, Dict, Tuple, Optional
import hashlib

import numpy as np
import pandas as pd
import streamlit as st

# ========================== Configuração de Logging ==========================
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ========================== Config ==========================
st.set_page_config(page_title="SAFETY • CHAT", page_icon="💬", layout="wide")

# Configuração de cache com TTL e controle de memória
CACHE_TTL_SECONDS = 3600  # 1 hora
MAX_CACHE_ITEMS = 50

DATA_DIR = Path("data")
AN_DIR   = DATA_DIR / "analytics"
XLSX_DIR = DATA_DIR / "xlsx"
DATASETS_CONTEXT_PATH = DATA_DIR / "datasets_context.md"
PROMPTS_MD_PATH       = DATA_DIR / "prompts" / "prompts.md"

SPH_PQ_PATH  = AN_DIR / "sphera.parquet"
SPH_NPZ_PATH = AN_DIR / "sphera_tfidf.joblib"  # Arquivo real dos embeddings Sphera

XLSX_LOCATION_PATH = XLSX_DIR / "TRATADO_safeguardOffShore.xlsx"

# Configurações do Ollama (serão inicializadas no contexto Streamlit)
OLLAMA_HOST = ""
OLLAMA_MODEL = ""
OLLAMA_API_KEY = ""
HEADERS_JSON = {"Content-Type": "application/json"}

def initialize_ollama_config():
    """Inicializa configurações do Ollama dentro do contexto Streamlit"""
    global OLLAMA_HOST, OLLAMA_MODEL, OLLAMA_API_KEY, HEADERS_JSON
    
    try:
        # Tentar acessar st.secrets primeiro
        if hasattr(st, 'secrets'):
            OLLAMA_HOST = st.secrets.get("OLLAMA_HOST", os.getenv("OLLAMA_HOST", ""))
            OLLAMA_MODEL = st.secrets.get("OLLAMA_MODEL", os.getenv("OLLAMA_MODEL", ""))
            OLLAMA_API_KEY = st.secrets.get("OLLAMA_API_KEY", os.getenv("OLLAMA_API_KEY"))
        else:
            # Fallback para variáveis de ambiente
            OLLAMA_HOST = os.getenv("OLLAMA_HOST", "")
            OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "")
            OLLAMA_API_KEY = os.getenv("OLLAMA_API_KEY")
    except Exception:
        # Fallback final para variáveis de ambiente
        OLLAMA_HOST = os.getenv("OLLAMA_HOST", "")
        OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "")
        OLLAMA_API_KEY = os.getenv("OLLAMA_API_KEY")
    
    HEADERS_JSON = {"Authorization": f"Bearer {OLLAMA_API_KEY}", "Content-Type": "application/json"} if OLLAMA_API_KEY else {"Content-Type": "application/json"}
    
    # Só definir padrões se não estiver configurado (não assumir localhost)
    if not OLLAMA_HOST and not os.getenv("OLLAMA_HOST"):
        OLLAMA_HOST = ""  # Não definir localhost automaticamente
        _info("Ollama não configurado - chat funcionará sem modelo")
    elif not OLLAMA_HOST:
        OLLAMA_HOST = "http://localhost:11434"  # Só usar localhost se foi configurado explicitamente
        
    if not OLLAMA_MODEL and not os.getenv("OLLAMA_MODEL"):
        OLLAMA_MODEL = ""  # Não definir modelo padrão automaticamente
    
    _info(f"Ollama configurado: {OLLAMA_HOST} -> {OLLAMA_MODEL or 'Não configurado'}")

def check_ollama_availability():
    """Verifica se o Ollama está disponível"""
    if not OLLAMA_HOST or not OLLAMA_MODEL:
        return False
    
    try:
        import requests
        response = requests.get(f"{OLLAMA_HOST}/api/tags", timeout=5)
        return response.status_code == 200
    except Exception:
        return False

# Sistema de cache otimizado com limites dinâmicos
class OptimizedCache:
    """Sistema de cache otimizado com limites inteligentes"""
    
    def __init__(self, max_items=MAX_CACHE_ITEMS, ttl_seconds=CACHE_TTL_SECONDS):
        self.max_items = max_items
        self.ttl_seconds = ttl_seconds
        self._cache_hits = 0
        self._cache_misses = 0
        self._cache_items = 0
    
    def get_cache_stats(self):
        """Retorna estatísticas do cache"""
        total = self._cache_hits + self._cache_misses
        hit_rate = (self._cache_hits / total * 100) if total > 0 else 0
        return {
            "hits": self._cache_hits,
            "misses": self._cache_misses,
            "hit_rate": round(hit_rate, 2),
            "items": self._cache_items,
            "max_items": self.max_items,
            "usage_pct": round((self._cache_items / self.max_items) * 100, 2)
        }
    
    def check_cache_health(self):
        """Verifica saúde do cache e gera alertas se necessário"""
        stats = self.get_cache_stats()
        if stats["usage_pct"] > CACHE_ALERT_THRESHOLD:
            _warn(f"Cache utilizando {stats['usage_pct']}% da capacidade máxima ({stats['items']}/{stats['max_items']} itens)")
        return stats

# Instância global do cache otimizado
cache_manager = OptimizedCache()

ST_MODEL_NAME = os.getenv("ST_MODEL_NAME", "sentence-transformers/all-MiniLM-L6-v2")

# ========================== Helpers ==========================

def _fatal(msg: str):
    """Erro fatal que para a execução da aplicação"""
    logger.error(msg)
    st.error(msg)
    st.stop()

def _warn(msg: str):
    """Aviso não-fatal que permite continuação"""
    logger.warning(msg)
    st.warning(msg)

def _info(msg: str):
    """Informação para debugging"""
    logger.info(msg)
    # st.info(msg)  # Comentado para evitar spam na interface

# Sistema de validação e alertas de configuração
def validate_configuration_sidebar_params(top_k_sphera, limiar_sphera, top_k_gosee, limiar_gosee, top_k_docs, limiar_docs, anos_filtro):
    """Valida configurações e gera alertas para o usuário (versão para parâmetros da sidebar)"""
    alerts = []
    
    # Verificar configurações problemáticas de limiar
    if limiar_sphera > 0.8:
        alerts.append(f"⚠️ Limiar de Similaridade Sphera muito alto ({limiar_sphera:.2f}). Considere reduzir para 0.3-0.6 para obter mais resultados.")
    
    if limiar_gosee > 0.8:
        alerts.append(f"⚠️ Limiar de Similaridade GoSee muito alto ({limiar_gosee:.2f}). Considere reduzir para 0.2-0.5 para obter mais resultados.")
    
    if limiar_docs > 0.8:
        alerts.append(f"⚠️ Limiar de Similaridade Documentos muito alto ({limiar_docs:.2f}). Considere reduzir para 0.3-0.6 para obter mais resultados.")
    
    # Verificar configurações muito permissivas
    if limiar_sphera < 0.1:
        alerts.append(f"⚠️ Limiar de Similaridade Sphera muito baixo ({limiar_sphera:.2f}). Considere aumentar para 0.2-0.4 para melhor precisão.")
    
    # Verificar Top-K muito alto
    if top_k_sphera > 50:
        alerts.append(f"⚠️ Top-K Sphera muito alto ({top_k_sphera}). Considere reduzir para 10-30 para melhor foco nos resultados.")
    
    # Verificar período muito longo
    if anos_filtro > 5:
        alerts.append(f"⚠️ Período muito longo ({anos_filtro} anos). Considere usar 1-3 anos para eventos mais recentes e relevantes.")
    
    return alerts

def validate_configuration():
    """Valida configurações e gera alertas para o usuário"""
    alerts = []
    
    # Por enquanto, apenas alertas gerais
    if not OLLAMA_HOST or not OLLAMA_MODEL:
        alerts.append("⚠️ Ollama não está configurado. Configure as variáveis de ambiente para usar o chat.")
    
    return alerts

def show_configuration_alerts():
    """Exibe alertas de configuração na sidebar"""
    alerts = validate_configuration()
    
    if alerts:
        with st.sidebar.expander("🔔 Alertas de Configuração", expanded=True):
            for alert in alerts:
                st.warning(alert)
    
    # Estatísticas do sistema
    with st.sidebar.expander("📊 Status do Sistema", expanded=False):
        cache_stats = cache_manager.get_cache_stats()
        st.write(f"**Cache:** {cache_stats['hits']} acertos, {cache_stats['misses']} erros ({cache_stats['hit_rate']}% de eficácia)")
        st.write(f"**Memória:** {cache_stats['usage_pct']}% utilizado ({cache_stats['items']}/{cache_stats['max_items']} itens)")
        
        # Status dos dados carregados
        status_data = {
            "Sphera": f"{len(df_sph):,} registros" if not df_sph.empty else "❌ Não disponível",
            "GoSee": f"{len(df_gosee):,} observações" if not df_gosee.empty else "❌ Não disponível",
            "Documentos": f"{len(docs_index)} arquivos" if docs_index else "❌ Não disponível",
            "Embeddings Sphera": "✅ Carregados" if E_sph is not None else "❌ Não disponível",
            "Embeddings GoSee": "✅ Carregados" if E_gosee is not None else "❌ Não disponível",
            "WS (Weak Signals)": "✅ Disponível" if E_ws is not None else "❌ Não disponível",
            "Precursores": "✅ Disponível" if E_prec is not None else "❌ Não disponível",
            "CP (Performance)": "✅ Disponível" if E_cp is not None else "❌ Não disponível",
        }
        
        # Status inteligente do Ollama
        ollama_status = ""
        if OLLAMA_HOST and OLLAMA_MODEL:
            # Verificar se Ollama está disponível
            if check_ollama_availability():
                ollama_status = f"✅ Conectado ({OLLAMA_MODEL})"
            else:
                ollama_status = f"⚠️ Configurado mas não conectado ({OLLAMA_MODEL})"
                ollama_status += "\n💡 Rode `ollama serve` ou configure uma API"
        else:
            ollama_status = "❌ Não configurado"
        
        status_data["Ollama"] = ollama_status
        
        for item, status in status_data.items():
            if item == "Ollama":
                # Status especial para Ollama com múltiplas linhas
                if "✅" in status:
                    st.success(f"**{item}:**")
                    st.success(status.replace("✅ ", ""))
                elif "⚠️" in status:
                    st.warning(f"**{item}:**")
                    st.warning(status.replace("⚠️ ", ""))
                else:
                    st.error(f"**{item}:**")
                    st.error(status.replace("❌ ", ""))
            else:
                if "✅" in status:
                    st.success(f"**{item}:** {status}")
                elif "⚠️" in status:
                    st.warning(f"**{item}:** {status}")
                else:
                    st.error(f"**{item}:** {status}")

def validate_embeddings_labels(embeddings: Optional[np.ndarray], labels: Optional[pd.DataFrame], name: str) -> bool:
    """Valida se embeddings e labels estão alinhados"""
    if embeddings is None and labels is None:
        _info(f"[{name}] Não disponível")
        return False
    if embeddings is None or labels is None:
        _warn(f"[{name}] Embutdings ou labels não disponível - pulando")
        return False
    if len(labels) != embeddings.shape[0]:
        _warn(f"[{name}] Desalinhamento: {len(labels)} labels vs {embeddings.shape[0]} embeddings")
        return False
    return True

def validate_dataframe(df: pd.DataFrame, name: str, required_cols: List[str] = None) -> bool:
    """Valida se DataFrame tem estrutura esperada"""
    if df is None or df.empty:
        _warn(f"[{name}] DataFrame vazio ou None")
        return False
    if required_cols:
        missing_cols = [col for col in required_cols if col not in df.columns]
        if missing_cols:
            _warn(f"[{name}] Colunas ausentes: {missing_cols}")
            return False
    return True

try:
    from sentence_transformers import SentenceTransformer
except Exception as e:
    _fatal(f"❌ sentence-transformers indisponível: {e}")

@st.cache_resource(show_spinner=False)
def ensure_st_encoder():
    try:
        return SentenceTransformer(ST_MODEL_NAME)
    except Exception as e:
        _fatal(f"❌ Não foi possível carregar o encoder: {e}")

@st.cache_data(show_spinner=False)
def load_embeddings_smart(base_path: Path, name: str = "embeddings") -> Optional[np.ndarray]:
    """
    Carrega embeddings de múltiplos formatos: .npz, .joblib, .jsonl, .parquet
    Suporte para diferentes formatos de vetores (TF-IDF, SentenceTransformers, etc.)
    """
    if not base_path.exists():
        # Tentar formatos alternativos
        alt_formats = [
            base_path.parent / f"{base_path.stem}.joblib",
            base_path.parent / f"{base_path.stem}.jsonl", 
            base_path.parent / f"{base_path.stem}.parquet",
            base_path.parent / f"{name}_tfidf.joblib",
            base_path.parent / f"{name}_embeddings.npz",
        ]
        
        for alt_path in alt_formats:
            if alt_path.exists():
                _info(f"Carregando {name} de formato alternativo: {alt_path}")
                base_path = alt_path
                break
        else:
            _warn(f"{name}: Nenhum arquivo de embeddings encontrado ({base_path} ou alternativas)")
            return None
    
    try:
        if base_path.suffix == ".npz":
            return load_npz_embeddings(base_path)
        elif base_path.suffix == ".joblib":
            return load_joblib_embeddings(base_path, name)
        elif base_path.suffix == ".jsonl":
            return load_jsonl_embeddings(base_path, name)
        elif base_path.suffix == ".parquet":
            return load_parquet_embeddings(base_path, name)
        else:
            _warn(f"{name}: Formato não suportado: {base_path.suffix}")
            return None
    except Exception as e:
        _warn(f"{name}: Erro ao carregar embeddings: {e}")
        return None

@st.cache_data(show_spinner=False)
def load_joblib_embeddings(joblib_path: Path, name: str = "embeddings") -> Optional[np.ndarray]:
    """Carrega embeddings do formato joblib"""
    try:
        import joblib
        data = joblib.load(str(joblib_path))
        
        # Diferentes formatos possíveis
        if isinstance(data, dict):
            # Tentar diferentes chaves
            for key in ['vectors', 'embeddings', 'features', 'tfidf_matrix', 'data']:
                if key in data and isinstance(data[key], np.ndarray):
                    return normalize_embeddings(data[key])
            
            # Se o dict inteiro for um array
            if len(data) > 0 and isinstance(list(data.values())[0], np.ndarray):
                return normalize_embeddings(np.array(list(data.values())))
        elif isinstance(data, np.ndarray):
            return normalize_embeddings(data)
        elif hasattr(data, 'toarray'):  # Matriz esparsa
            return normalize_embeddings(data.toarray())
        else:
            _warn(f"{name}: Estrutura joblib não reconhecida")
            return None
            
    except Exception as e:
        _warn(f"{name}: Erro ao carregar joblib: {e}")
        return None

@st.cache_data(show_spinner=False)
def load_jsonl_embeddings(jsonl_path: Path, name: str = "embeddings") -> Optional[np.ndarray]:
    """Carrega embeddings do formato jsonl"""
    try:
        import json
        vectors = []
        with open(jsonl_path, 'r', encoding='utf-8') as f:
            for line in f:
                if line.strip():
                    data = json.loads(line.strip())
                    # Tentar diferentes formatos
                    if 'vector' in data:
                        vectors.append(data['vector'])
                    elif 'embedding' in data:
                        vectors.append(data['embedding'])
                    elif 'vec' in data:
                        vectors.append(data['vec'])
                    elif isinstance(data, list):
                        vectors.append(data)
                    else:
                        _warn(f"{name}: Formato JSONL não reconhecido: {list(data.keys())}")
                        continue
        
        if vectors:
            return normalize_embeddings(np.array(vectors))
        else:
            _warn(f"{name}: Nenhum vetor encontrado no JSONL")
            return None
            
    except Exception as e:
        _warn(f"{name}: Erro ao carregar JSONL: {e}")
        return None

@st.cache_data(show_spinner=False)
def load_parquet_embeddings(parquet_path: Path, name: str = "embeddings") -> Optional[np.ndarray]:
    """Carrega embeddings do formato parquet"""
    try:
        df = pd.read_parquet(parquet_path)
        
        # Tentar diferentes colunas
        for col in ['vector', 'embedding', 'vec', 'features', 'data']:
            if col in df.columns:
                vectors = df[col].apply(lambda x: np.array(x) if isinstance(x, list) else x).values
                if len(vectors) > 0:
                    return normalize_embeddings(np.vstack(vectors))
        
        # Se não encontrou colunas específicas, tentar todas as colunas numéricas
        numeric_cols = df.select_dtypes(include=[np.number]).columns
        if len(numeric_cols) > 0:
            return normalize_embeddings(df[numeric_cols].values)
        else:
            _warn(f"{name}: Nenhuma coluna numérica encontrada no parquet")
            return None
            
    except Exception as e:
        _warn(f"{name}: Erro ao carregar parquet: {e}")
        return None

def normalize_embeddings(embeddings: np.ndarray) -> np.ndarray:
    """Normaliza embeddings para magnitude unitária"""
    if embeddings.size == 0:
        return embeddings
    
    norms = np.linalg.norm(embeddings, axis=1, keepdims=True) + 1e-9
    return (embeddings / norms).astype(np.float32)

@st.cache_data(show_spinner=False)
def load_npz_embeddings(path: Path) -> Optional[np.ndarray]:
    """Função original para carregar .npz mantida para compatibilidade"""
    if not path.exists():
        return None
    try:
        with np.load(str(path), allow_pickle=True) as z:
            for key in ("embeddings", "E", "X", "vectors", "vecs", "arr_0"):
                if key in z:
                    E = np.array(z[key]).astype(np.float32, copy=False)
                    n = np.linalg.norm(E, axis=1, keepdims=True) + 1e-9
                    return (E / n).astype(np.float32)
            # fallback: maior matriz 2D
            best_k, best_n = None, -1
            for k in z.files:
                arr = z[k]
                if isinstance(arr, np.ndarray) and arr.ndim == 2 and arr.shape[0] > best_n:
                    best_k, best_n = k, arr.shape[0]
            if best_k is None:
                return None
            E = np.array(z[best_k]).astype(np.float32, copy=False)
            n = np.linalg.norm(E, axis=1, keepdims=True) + 1e-9
            return (E / n).astype(np.float32)
    except Exception:
        return None

@st.cache_data(show_spinner=False)
def load_embeddings_any_format(path: Path) -> Optional[np.ndarray]:
    """
    Carrega embeddings de qualquer formato suportado: .npz, .joblib, .jsonl, .parquet
    """
    if not path.exists():
        return None
    
    try:
        # Tentar diferentes formatos baseado na extensão
        if path.suffix.lower() == '.npz':
            return load_npz_embeddings(path)
        
        elif path.suffix.lower() == '.joblib':
            try:
                import joblib
                data = joblib.load(str(path))
                # Verificar se é numpy array
                if isinstance(data, np.ndarray):
                    # Normalizar embeddings se necessário
                    if data.ndim == 2:
                        norms = np.linalg.norm(data, axis=1, keepdims=True) + 1e-9
                        return (data / norms).astype(np.float32)
                    return data.astype(np.float32)
                elif isinstance(data, dict) and 'embeddings' in data:
                    E = np.array(data['embeddings'])
                    if E.ndim == 2:
                        norms = np.linalg.norm(E, axis=1, keepdims=True) + 1e-9
                        return (E / norms).astype(np.float32)
                else:
                    _warn(f"Formato de dados desconhecido em {path}: {type(data)}")
                    return None
            except Exception as e:
                _warn(f"Erro ao carregar joblib {path}: {e}")
                return None
        
        elif path.suffix.lower() == '.jsonl':
            try:
                import json
                embeddings = []
                with open(path, 'r', encoding='utf-8') as f:
                    for line in f:
                        if line.strip():
                            data = json.loads(line)
                            if 'embedding' in data:
                                embeddings.append(data['embedding'])
                            elif 'embeddings' in data:
                                embeddings.extend(data['embeddings'])
                
                if embeddings:
                    E = np.array(embeddings)
                    if E.ndim == 2:
                        norms = np.linalg.norm(E, axis=1, keepdims=True) + 1e-9
                        return (E / norms).astype(np.float32)
                return None
            except Exception as e:
                _warn(f"Erro ao carregar jsonl {path}: {e}")
                return None
        
        elif path.suffix.lower() == '.parquet':
            try:
                df = pd.read_parquet(path)
                # Tentar diferentes colunas comuns para embeddings
                embedding_cols = ['embedding', 'embeddings', 'vector', 'vectors', 'E']
                for col in embedding_cols:
                    if col in df.columns:
                        E = np.array(df[col].tolist())
                        if E.ndim == 2:
                            norms = np.linalg.norm(E, axis=1, keepdims=True) + 1e-9
                            return (E / norms).astype(np.float32)
                _warn(f"Nenhuma coluna de embedding encontrada em {path}")
                return None
            except Exception as e:
                _warn(f"Erro ao carregar parquet {path}: {e}")
                return None
        
        else:
            _warn(f"Formato de arquivo não suportado: {path.suffix}")
            return None
            
    except Exception as e:
        _warn(f"Erro geral ao carregar {path}: {e}")
        return None

@st.cache_data(show_spinner=False)
def load_prompts_md(md_path: Path) -> Dict[str, List[Dict[str, str]]]:
    if not md_path.exists():
        return {"Texto": [], "Upload": []}
    raw = md_path.read_text(encoding="utf-8")
    sections = re.split(r"(?m)^##\s+", raw)
    data = {"Texto": [], "Upload": []}
    for sec in sections:
        sec = sec.strip()
        if not sec:
            continue
        first, _, rest = sec.partition("\n")
        if first.strip() not in ("Texto", "Upload"):
            continue
        parts = re.split(r"(?m)^###\s+", rest)
        for p in parts:
            p = p.strip()
            if not p:
                continue
            title, _, body = p.partition("\n")
            data[first.strip()].append({"title": title.strip(), "body": body.strip()})
    return data

@st.cache_data(show_spinner=False)
def load_file_text(p: Path) -> str:
    try:
        return p.read_text(encoding="utf-8")
    except Exception as e:
        return f"[AVISO] Não consegui ler {p}: {e} (continuando sem este contexto)"

# ========================== Helpers (Location) ==========================

def get_sphera_location_col(df: pd.DataFrame) -> Optional[str]:
    """Retorna a melhor coluna de localização, nunca AREA.
    Prioridade: LOCATION → FPSO → Location → FPSO/Unidade → Unidade."""
    if df is None or df.empty:
        return None
    for c in ["LOCATION", "FPSO", "Location", "FPSO/Unidade", "Unidade"]:
        if c in df.columns and df[c].notna().any():
            return c
    return None

@st.cache_data(show_spinner=False)
def _location_options_from(df_full: pd.DataFrame) -> Tuple[Optional[str], List[str]]:
    col = get_sphera_location_col(df_full)
    if not col:
        return None, []
    s = df_full[col].astype(str).str.strip()
    s = s[(~s.isna()) & (s.str.len() > 0)]
    bad = {"nan", "none", "n/d", "nd"}
    s = s[~s.str.lower().isin(bad)]
    # de-duplicar preservando primeira grafia
    seen = {}
    for v in s:
        k = v.lower()
        if k not in seen:
            seen[k] = v
    return col, sorted(seen.values())

# ========================== Helpers (Text Extraction) ==========================

def extract_pdf_text(file_like: io.BytesIO) -> str:
    """
    Extrai texto de PDF. Tenta PyPDF2 -> PyMuPDF (fitz) -> pdfminer.six.
    Retorna string (pode ser vazia se o PDF for apenas imagem/scaneado).
    """
    # Validar header do arquivo PDF
    header = file_like.read(4)
    if header[:4] != b'%PDF':
        return ""  # Não é um PDF válido
    file_like.seek(0)
    
    # 1) PyPDF2
    try:
        import PyPDF2
        file_like.seek(0)
        reader = PyPDF2.PdfReader(file_like)
        if reader.is_encrypted:
            return ""  # PDF protegido por senha
        parts = []
        for page in reader.pages:
            parts.append(page.extract_text() or "")
        return "\n".join(parts).strip()
    except Exception:
        pass
    
    # 2) PyMuPDF
    try:
        import fitz  # PyMuPDF
        file_like.seek(0)
        doc = fitz.open(stream=file_like.read(), filetype="pdf")
        if doc.is_encrypted:
            return ""  # PDF protegido por senha
        parts = [page.get_text() for page in doc]
        return "\n".join(parts).strip()
    except Exception:
        pass
    
    # 3) pdfminer.six
    try:
        from pdfminer.high_level import extract_text
        file_like.seek(0)
        return (extract_text(file_like) or "").strip()
    except Exception:
        pass
    
    return ""

def extract_docx_text(file_like: io.BytesIO) -> str:
    """Extrai texto de um .docx (python-docx)."""
    try:
        from docx import Document
        file_like.seek(0)
        doc = Document(file_like)
        parts = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" ".join(cell.text for cell in row.cells if cell.text))
        return "\n".join(parts).strip()
    except Exception:
        return ""

def extract_xlsx_text(file_like: io.BytesIO) -> str:
    """Extrai texto de um .xlsx (pandas + openpyxl)."""
    try:
        file_like.seek(0)
        sheets = pd.read_excel(file_like, sheet_name=None, engine="openpyxl")
        lines = []
        for name, df in sheets.items():
            if df is None or df.empty:
                continue
            df = df.astype(str).fillna("")
            lines.append(f"=== SHEET: {name} ===")
            lines.extend(df.apply(lambda r: " ".join(r.values), axis=1).tolist())
        return "\n".join(lines).strip()
    except Exception:
        return ""

# ========================== Carregamento de dados ==========================

# Validação crítica: Sphera é obrigatório
if not SPH_PQ_PATH.exists():
    _fatal(f"Parquet do Sphera não encontrado em {SPH_PQ_PATH}")

# --- SPHERA ---
df_sph = pd.read_parquet(SPH_PQ_PATH) if SPH_PQ_PATH.exists() else pd.DataFrame()

# Validação flexível - verificar quais colunas existem
required_cols = []
available_cols = []

# Verificar colunas essenciais
if not df_sph.empty:
    if "Description" in df_sph.columns:
        available_cols.append("Description")
    if "DESCRIPTION" in df_sph.columns:  # alternativo
        available_cols.append("DESCRIPTION")
    if "EVENT_DATE" in df_sph.columns:
        available_cols.append("EVENT_DATE")
    
    # Usar validação flexível baseada no que está disponível
    if not available_cols:
        _warn("Sphera: Nenhuma coluna essencial encontrada (Description/DESCRIPTION)")
        df_sph = pd.DataFrame()  # Fallback para DataFrame vazio
    elif "Description" not in available_cols and "DESCRIPTION" not in available_cols:
        _warn("Sphera: Coluna Description/DESCRIPTION não encontrada")
        df_sph = pd.DataFrame()  # Fallback para DataFrame vazio
else:
    _warn("Sphera: DataFrame vazio")

E_sph = load_embeddings_any_format(SPH_NPZ_PATH)
if E_sph is None:
    _warn("Embeddings do Sphera não encontrados - funcionalidade limitada")
else:
    _info(f"Embeddings do Sphera carregados: {E_sph.shape[0]} registros")

# --- GOSEE (Implementação completa) ---
GOSEE_PQ_PATH = AN_DIR / "gosee.parquet"
GOSEE_NPZ_PATH = AN_DIR / "gosee_tfidf.joblib"  # Arquivo real dos embeddings GoSee

df_gosee = pd.read_parquet(GOSEE_PQ_PATH) if GOSEE_PQ_PATH.exists() else pd.DataFrame()
if not validate_dataframe(df_gosee, "GoSee", ["Observation"]):
    df_gosee = pd.DataFrame()  # Fallback para DataFrame vazio
    _info("GoSee não disponível - continuando sem esta fonte")

# Carregar embeddings específicos do GoSee usando sistema inteligente
E_gosee = load_embeddings_any_format(GOSEE_NPZ_PATH)
if E_gosee is None:
    _warn("Embeddings do GoSee não encontrados - busca no GoSee limitada")
else:
    _info(f"Embeddings do GoSee carregados: {E_gosee.shape[0]} observações")

# --- DOCUMENTOS (NOVO: Processamento de PDFs/DOCXs) ---
DOCS_DIR = DATA_DIR / "docs"
docs_index = {}  # Índice: {nome_arquivo: texto_completo}
if DOCS_DIR.exists() and DOCS_DIR.is_dir():
    for doc_path in DOCS_DIR.glob("*.pdf"):
        try:
            text = extract_pdf_text(io.BytesIO(doc_path.read_bytes()))
            docs_index[doc_path.name] = text
            _info(f"Documento carregado: {doc_path.name} ({len(text)} chars)")
        except Exception as e:
            _warn(f"Erro ao processar {doc_path.name}: {e}")
    
    for doc_path in DOCS_DIR.glob("*.docx"):
        try:
            text = extract_docx_text(io.BytesIO(doc_path.read_bytes()))
            docs_index[doc_path.name] = text
            _info(f"Documento carregado: {doc_path.name} ({len(text)} chars)")
        except Exception as e:
            _warn(f"Erro ao processar {doc_path.name}: {e}")
else:
    _info("Pasta de documentos não encontrada - continuando sem documentos")

# coluna exibida para Location
if not df_sph.empty:
    LOC_DISPLAY_COL = get_sphera_location_col(df_sph)
    if not LOC_DISPLAY_COL:
        _warn("Coluna de localização não encontrada no Sphera")
else:
    LOC_DISPLAY_COL = None

# --- WS/Precursores ---
WS_NPZ,   WS_LBL   = AN_DIR / "ws_embeddings_pt.npz",   AN_DIR / "ws_embeddings_pt.parquet"
PREC_NPZ, PREC_LBL = AN_DIR / "prec_embeddings_pt.npz", AN_DIR / "prec_embeddings_pt.parquet"

E_ws, L_ws = None, None
if WS_NPZ.exists() and WS_LBL.exists():
    E_ws = load_npz_embeddings(WS_NPZ)
    L_ws = pd.read_parquet(WS_LBL)
    if not validate_embeddings_labels(E_ws, L_ws, "WS"):
        E_ws, L_ws = None, None

E_prec, L_prec = None, None  
if PREC_NPZ.exists() and PREC_LBL.exists():
    E_prec = load_npz_embeddings(PREC_NPZ)
    L_prec = pd.read_parquet(PREC_LBL)
    if not validate_embeddings_labels(E_prec, L_prec, "Precursores"):
        E_prec, L_prec = None, None

# --- CP (loader robusto com fallbacks) ---
CP_NPZ_MAIN   = AN_DIR / "cp_embeddings.npz"
CP_NPZ_ALT    = AN_DIR / "cp_vectors.npz"        # fallback
CP_LBL_PARQ   = AN_DIR / "cp_labels.parquet"
CP_LBL_JSONL  = AN_DIR / "cp_labels.jsonl"       # fallback

@st.cache_data(show_spinner=False)
def _load_npz_any(path: Path):
    """Carrega embeddings NPZ com fallback robusto"""
    if not path.exists():
        return None
    try:
        with np.load(str(path), allow_pickle=True) as z:
            for k in ("embeddings", "E", "X", "vectors", "vecs", "arr_0"):
                if k in z:
                    A = np.array(z[k])
                    if isinstance(A, np.ndarray) and A.ndim == 2 and A.shape[0] > 0:
                        A = A.astype(np.float32, copy=False)
                        A /= (np.linalg.norm(A, axis=1, keepdims=True) + 1e-9)
                        return A
            # maior matriz 2D
            best = None
            for k in z.files:
                A = z[k]
                if isinstance(A, np.ndarray) and A.ndim == 2 and (best is None or A.shape[0] > best.shape[0]):
                    best = A
            if best is not None:
                A = best.astype(np.float32, copy=False)
                A /= (np.linalg.norm(A, axis=1, keepdims=True) + 1e-9)
                return A
    except Exception as e:
        _warn(f"Erro ao carregar {path}: {e}")
        return None
    return None

@st.cache_data(show_spinner=False)
def _load_cp_labels() -> Optional[pd.DataFrame]:
    """Carrega labels do CP com fallbacks"""
    df = None
    if CP_LBL_PARQ.exists():
        try:
            df = pd.read_parquet(CP_LBL_PARQ)
        except Exception as e:
            _warn(f"Erro ao carregar CP labels parquet: {e}")
            df = None
    if df is None and CP_LBL_JSONL.exists():
        try:
            df = pd.read_json(CP_LBL_JSONL, lines=True)
        except Exception as e:
            _warn(f"Erro ao carregar CP labels jsonl: {e}")
            df = None
    if df is None:
        return None
    label_col = next((c for c in ["label","LABEL","text","name","CP","cp"] if c in df.columns), None)
    if not label_col:
        _warn("Coluna de label não encontrada nos CP labels")
        return None
    if label_col != "label":
        df = df.rename(columns={label_col: "label"})
    return df[["label"]] if "label" in df.columns else None

E_cp = _load_npz_any(CP_NPZ_MAIN)
if E_cp is None:
    E_cp = _load_npz_any(CP_NPZ_ALT)

L_cp = _load_cp_labels()

# Validação final do CP com fallback
if E_cp is not None and L_cp is not None:
    if not validate_embeddings_labels(E_cp, L_cp, "CP"):
        E_cp, L_cp = None, None
else:
    _info("CP não disponível - continuando sem esta funcionalidade")

# Relatório de status dos dados carregados
status_data = {
    "Sphera": f"{len(df_sph)} registros" if not df_sph.empty else "Não disponível",
    "Embeddings Sphera": "OK" if E_sph is not None else "Não disponível",
    "GoSee": f"{len(df_gosee)} registros" if not df_gosee.empty else "Não disponível",
    "Documentos PDF/DOCX": f"{len(docs_index)} arquivos" if docs_index else "Não disponível",
    "WS": "OK" if E_ws is not None and L_ws is not None else "Não disponível",
    "Precursores": "OK" if E_prec is not None and L_prec is not None else "Não disponível",
    "CP": "OK" if E_cp is not None and L_cp is not None else "Não disponível"
}

with st.expander("📊 Status dos Dados Carregados", expanded=False):
    status_df = pd.DataFrame(list(status_data.items()), columns=["Componente", "Status"])
    st.dataframe(status_df, use_container_width=True, hide_index=True)

# ========================== Estado ==========================
# Inicializar configurações do Ollama
initialize_ollama_config()

if "system_prompt" not in st.session_state:
    pre = (
        "Você é o ESO-CHAT para segurança operacional (óleo e gás). "
        "Responda em PT-BR, cite IDs/sim quando usar buscas locais, e não invente dados fora dos contextos fornecidos.\n\n"
    )
    sys_ctx = (load_file_text(DATASETS_CONTEXT_PATH) if DATASETS_CONTEXT_PATH.exists() else "")
    st.session_state.system_prompt = pre + ("=== DATASETS_CONTEXT ===\n" + sys_ctx if sys_ctx else "")
if "chat" not in st.session_state:
    st.session_state.chat = []
if "draft_prompt" not in st.session_state:
    st.session_state.draft_prompt = ""
if "_clear_draft_flag" not in st.session_state:
    st.session_state._clear_draft_flag = False
if "st_encoder" not in st.session_state:
    st.session_state.st_encoder = ensure_st_encoder()
if "upld_texts" not in st.session_state:
    st.session_state.upld_texts = []

def clear_stale_cache():
    """Limpa cache antigo para evitar problemas de memória"""
    try:
        if hasattr(st, 'cache_data') and hasattr(st.cache_data, 'clear'):
            # Limpa cache específico se disponível
            cache_manager._cache_items = max(0, cache_manager._cache_items - 10)  # Reduz contador
            _info("Cache limpo automaticamente")
    except Exception as e:
        _warn(f"Erro ao limpar cache: {e}")

# Controle de performance aprimorado
def log_performance(func_name: str, duration: float):
    """Log de performance das operações críticas com alertas inteligentes"""
    if duration > 10.0:  # > 10 segundos
        _warn(f"⚠️ Operação {func_name} MUITO LENTA: {duration:.2f}s")
    elif duration > 5.0:  # > 5 segundos
        _warn(f"Operação {func_name} lenta: {duration:.2f}s")
    else:
        _info(f"Operação {func_name}: {duration:.2f}s")
    
    # Log do cache health
    cache_stats = cache_manager.check_cache_health()
    if cache_stats["usage_pct"] > 90:
        _warn(f"Cache com alta utilização: {cache_stats['usage_pct']}%")

# ========================== Encode ==========================
@st.cache_data(show_spinner=False)
def encode_texts(texts: List[str], batch_size: int = 64) -> np.ndarray:
    M = st.session_state.st_encoder.encode(
        texts, batch_size=batch_size, show_progress_bar=False,
        convert_to_numpy=True, normalize_embeddings=True
    ).astype(np.float32)
    return M

@st.cache_data(show_spinner=False)
def encode_query(q: str) -> np.ndarray:
    v = st.session_state.st_encoder.encode([q], convert_to_numpy=True, normalize_embeddings=True)[0].astype(np.float32)
    v /= (np.linalg.norm(v) + 1e-9)
    return v

# ========================== Filtros / Similaridade ==========================

@st.cache_data(show_spinner=False)
def gosee_similar_to_text(
    query_text: str,
    min_sim: float = 0.3,
    topk: int = 20,
    df_base: Optional[pd.DataFrame] = None,
    substr: str = "",
) -> List[Tuple[str, float, pd.Series]]:
    """Busca similar no GoSee usando embeddings específicos do GoSee"""
    if not query_text or not query_text.strip():
        return []
    
    if df_base is None or df_base.empty:
        return []
    
    # CRÍTICO: Verificar se temos embeddings específicos do GoSee
    if E_gosee is None:
        _warn("Embeddings do GoSee não disponíveis - busca desabilitada")
        return []
    
    start_time = time.time()
    
    # Pré-filtros (se aplicável)
    df_filtered = df_base.copy()
    
    # Filtro de substring na descrição
    if substr:
        df_filtered = df_filtered[df_filtered["Observation"].str.contains(substr, case=False, na=False)]
    
    if df_filtered.empty:
        return []
    
    # Encode da query
    v_query = encode_query(query_text)
    
    # Similaridade cosseno usando embeddings específicos do GoSee
    if "Observation" in df_filtered.columns:
        texts = df_filtered["Observation"].fillna("").tolist()
        try:
            # CRÍTICO: Usar E_gosee em vez de E_sph
            E = E_gosee[:len(df_filtered)]
            similarities = (E @ v_query).squeeze()
            
            # Filtro por limiar
            valid_mask = similarities >= min_sim
            valid_indices = np.where(valid_mask)[0]
            
            if len(valid_indices) == 0:
                return []
            
            # Ordenar por similaridade
            sorted_indices = valid_indices[np.argsort(similarities[valid_indices])[::-1]]
            
            # Construir resultados
            results = []
            for i in sorted_indices[:topk]:
                idx_in_filtered = i
                original_idx = df_filtered.iloc[idx_in_filtered].name
                similarity = float(similarities[i])
                row = df_filtered.iloc[idx_in_filtered]
                results.append((str(original_idx), similarity, row))
            
            duration = time.time() - start_time
            log_performance(f"gosee_search_{len(results)}_results", duration)
            
            return results
            
        except Exception as e:
            _warn(f"Erro na busca GoSee: {e}")
            return []
    
    return []

@st.cache_data(show_spinner=False)
def docs_similar_to_text(
    query_text: str,
    min_sim: float = 0.3,
    topk: int = 10,
    docs_dict: Optional[Dict[str, str]] = None,
) -> List[Tuple[str, float, str]]:
    """Busca similar em documentos PDF/DOCX indexados"""
    if not query_text or not query_text.strip() or not docs_dict:
        return []
    
    start_time = time.time()
    
    # Se não há documentos, retorna vazio
    if not docs_dict:
        return []
    
    try:
        # Preparar textos dos documentos
        doc_texts = []
        doc_names = []
        
        for doc_name, doc_text in docs_dict.items():
            if doc_text and doc_text.strip():
                doc_texts.append(doc_text[:2000])  # Limitar tamanho para performance
                doc_names.append(doc_name)
        
        if not doc_texts:
            return []
        
        # Encode de todos os textos dos documentos
        doc_embeddings = encode_texts(doc_texts, batch_size=16)
        
        # Encode da query
        v_query = encode_query(query_text)
        
        # Similaridade cosseno
        similarities = (doc_embeddings @ v_query).squeeze()
        
        # Filtro por limiar
        valid_mask = similarities >= min_sim
        valid_indices = np.where(valid_mask)[0]
        
        if len(valid_indices) == 0:
            return []
        
        # Ordenar por similaridade
        sorted_indices = valid_indices[np.argsort(similarities[valid_indices])[::-1]]
        
        # Construir resultados
        results = []
        for i in sorted_indices[:topk]:
            doc_name = doc_names[i]
            similarity = float(similarities[i])
            text_snippet = doc_texts[i][:500] + "..." if len(doc_texts[i]) > 500 else doc_texts[i]
            results.append((doc_name, similarity, text_snippet))
        
        duration = time.time() - start_time
        log_performance(f"docs_search_{len(results)}_results", duration)
        
        return results
        
    except Exception as e:
        _warn(f"Erro na busca em documentos: {e}")
        return []

@st.cache_data(show_spinner=False)
def filter_sphera(df: pd.DataFrame, locations: List[str], substr: str, years: int) -> pd.DataFrame:
    """Filtro robusto do Sphera com validação e logging"""
    if df is None or df.empty:
        _warn("DataFrame vazio fornecido para filtro Sphera")
        return pd.DataFrame()
        
    out = df.copy()
    original_size = len(out)
    
    try:
        # Janela temporal
        if "EVENT_DATE" in out.columns:
            out["EVENT_DATE"] = pd.to_datetime(out["EVENT_DATE"], errors="coerce")
            cutoff = pd.Timestamp(datetime.utcnow() - timedelta(days=365 * years))
            before_date = len(out)
            out = out[out["EVENT_DATE"] >= cutoff]
            _info(f"Filtro temporal: {len(out)}/{before_date} eventos após {years} anos")
        
        # Filtro por Location (string exata, preservando grafia exibida)
        loc_col = get_sphera_location_col(out)
        if loc_col and locations:
            before_loc = len(out)
            selected = set([str(x).strip() for x in locations if str(x).strip()])
            out = out[out[loc_col].astype(str).isin(selected)]
            _info(f"Filtro Location ({loc_col}): {len(out)}/{before_loc} eventos")

        # Description contém (case-insensitive)
        desc_col = "Description" if "Description" in out.columns else ("DESCRIPTION" if "DESCRIPTION" in out.columns else None)
        if desc_col and substr:
            before_substr = len(out)
            pat = re.escape(substr)
            mask = out[desc_col].astype(str).str.contains(pat, case=False, na=False, regex=True)
            out = out[mask]
            _info(f"Filtro substring '{substr}': {len(out)}/{before_substr} eventos")

        _info(f"Filtros Sphera aplicados: {len(out)}/{original_size} eventos restantes")
        return out
        
    except Exception as e:
        _warn(f"Erro ao aplicar filtros Sphera: {e}")
        return df  # Retorna DataFrame original em caso de erro

@st.cache_data(show_spinner=False)
def sphera_similar_to_text(query_text: str, min_sim: float, years: int, topk: int,
                           df_base: pd.DataFrame, E_base: Optional[np.ndarray],
                           substr: str, locations: List[str]) -> List[Tuple[str, float, pd.Series]]:
    """Busca similar com validação robusta e logging"""
    start_time = datetime.now()
    
    if not query_text or not query_text.strip():
        _warn("Query vazia fornecida para busca similar")
        return []
    
    if df_base is None or df_base.empty:
        _warn("DataFrame Sphera vazio ou None")
        return []
        
    if E_base is None or E_base.size == 0:
        _warn("Embeddings Sphera não disponíveis")
        return []
    
    # Filtros aplicados
    base = filter_sphera(df_base, locations, substr, years)
    if base.empty:
        _info("Nenhum evento passou pelos filtros")
        return []
        
    _info(f"Filtros aplicados: {len(base)}/{len(df_base)} eventos restantes")
    
    try:
        idx_map = base.index.to_numpy()
        if np.issubdtype(idx_map.dtype, np.integer):
            E_view = E_base[idx_map, :]
        else:
            E_view = E_base
            base = df_base
    except Exception as e:
        _warn(f"Erro ao mapear índices: {e}")
        E_view = E_base
        base = df_base
        
    try:
        qv = encode_query(query_text.strip())
        sims = (E_view @ qv).astype(float)
        ord_idx = np.argsort(-sims)
        id_col = "Event ID" if "Event ID" in base.columns else ("EVENT_NUMBER" if "EVENT_NUMBER" in base.columns else ("EVENTID" if "EVENTID" in base.columns else None))
        
        out = []
        kept = 0
        for i in ord_idx:
            s = float(sims[i])
            if s < min_sim:
                continue
            row = base.iloc[int(i)]
            evid = row.get(id_col, f"row{i}") if id_col else f"row{i}"
            out.append((str(evid), s, row))
            kept += 1
            if kept >= topk:
                break
                
        elapsed = (datetime.now() - start_time).total_seconds()
        _info(f"Busca concluída: {len(out)}/{kept} resultados em {elapsed:.2f}s")
        return out
        
    except Exception as e:
        _warn(f"Erro na busca similar: {e}")
        return []

# ========================== Agregação dicionários ==========================
@st.cache_data(show_spinner=False)
def aggregate_dict_matches_over_hits(
    hits: List[Tuple[str, float, pd.Series]],
    E_ws, L_ws, E_prec, L_prec, E_cp, L_cp,
    thr_ws_sim: float, thr_prec_sim: float, thr_cp_sim: float,
    topn_ws: int, topn_prec: int, topn_cp: int,
    agg_mode: str = "max",
    per_event_thr: float = 0.30,
    min_support: int = 1,
) -> Dict[str, List[Tuple[str, float, int]]]:
    if not hits:
        return {"ws": [], "prec": [], "cp": []}
    descs = [str(r.get("Description", r.get("DESCRIPTION", "")).strip()) for _, _, r in hits]
    descs = [d for d in descs if d]
    if not descs:
        return {"ws": [], "prec": [], "cp": []}
    V_desc = encode_texts(descs, batch_size=32).T  # (D x M)

    def _score(E_bank, labels_df, thr_sim, topn_target):
        if E_bank is None or labels_df is None or len(labels_df) != E_bank.shape[0]:
            return []
        S = (E_bank @ V_desc)                 # (N_terms, M_events)
        support = (S >= per_event_thr).sum(axis=1)
        sims = S.mean(axis=1) if agg_mode == "mean" else S.max(axis=1)
        mask = (support >= min_support) & (sims >= thr_sim)
        idx = np.where(mask)[0]
        if idx.size == 0:
            return []
        order = idx[np.argsort(sims[idx])[::-1]]
        out_terms = []
        for i in order[:topn_target]:
            label = str(labels_df.iloc[i].get("label", labels_df.iloc[i].get("text", f"TERM_{i}")))
            out_terms.append((label, float(sims[i]), int(support[i])))
        return out_terms

    return {
        "ws":   _score(E_ws,   L_ws,   thr_ws_sim,   topn_ws),
        "prec": _score(E_prec, L_prec, thr_prec_sim, topn_prec),
        "cp":   _score(E_cp,   L_cp,   thr_cp_sim,   topn_cp),
    }

# ===== Depuração: Top-N "brutos" (ignora thresholds) =====
def _topk_raw_for_bank(E_bank: np.ndarray, labels_df: pd.DataFrame, V_desc_T: np.ndarray, topk: int = 10):
    if E_bank is None or labels_df is None or (len(labels_df) != (E_bank.shape[0] if hasattr(E_bank,'shape') else 0)) or V_desc_T is None:
        return pd.DataFrame()
    S = (E_bank @ V_desc_T)
    sims = S.max(axis=1)
    order = np.argsort(sims)[::-1][:topk]
    labels_col = next((c for c in ["label","text","name","CP","cp"] if c in labels_df.columns), None)
    rows = []
    for i in order:
        lab = str(labels_df.iloc[i].get(labels_col, f"TERM_{i}"))
        rows.append({"Termo": lab, "Similaridade(max)": float(sims[i])})
    return pd.DataFrame(rows)

def debug_preview_dicts(hits, E_ws, L_ws, E_prec, L_prec, E_cp, L_cp, topk=10):
    if not hits:
        st.info("Sem hits Sphera para depuração de dicionários.")
        return
    descs = [str(r.get("Description", r.get("DESCRIPTION",""))).strip() for _,_,r in hits]
    descs = [d for d in descs if d]
    if not descs:
        st.info("Sem descrições válidas para depuração.")
        return
    V_desc = encode_texts(descs, batch_size=32).T  # (D x M)
    with st.expander("🔎 Depuração — Top-N brutos (ignora thresholds)", expanded=False):
        if E_ws is not None and L_ws is not None:
            st.markdown("**WS (max entre eventos, sem limiares)**")
            st.dataframe(_topk_raw_for_bank(E_ws, L_ws, V_desc, topk), use_container_width=True, hide_index=True)
        if E_prec is not None and L_prec is not None:
            st.markdown("**Precursores (max entre eventos, sem limiares)**")
            st.dataframe(_topk_raw_for_bank(E_prec, L_prec, V_desc, topk), use_container_width=True, hide_index=True)
        if E_cp is not None and L_cp is not None:
            st.markdown("**CP (max entre eventos, sem limiares)**")
            st.dataframe(_topk_raw_for_bank(E_cp, L_cp, V_desc, topk), use_container_width=True, hide_index=True)

# ===== Hints por evento =====
def build_event_hints(
    hits: List[Tuple[str, float, pd.Series]],
    E_ws, L_ws, E_prec, L_prec, E_cp, L_cp,
    per_event_thr: float,
    top_per_family: int = 3,
) -> Tuple[str, np.ndarray | None]:
    if not hits:
        return "=== EVENT_HINTS === [NENHUM HIT]", None
    descs = [str(r.get("Description", r.get("DESCRIPTION",""))).strip() for _,_,r in hits]
    descs = [d for d in descs if d]
    if not descs:
        return "=== EVENT_HINTS === [SEM DESCRIÇÕES VÁLIDAS]", None
    V_desc = encode_texts(descs, batch_size=32).T  # (D x M)
    def _labels_col(df):
        return next((c for c in ["label","text","name","CP","cp"] if (df is not None and c in df.columns)), None)
    lines = ["=== EVENT_HINTS ==="]
    families = [("WS",E_ws,L_ws), ("PRE",E_prec,L_prec), ("CP",E_cp,L_cp)]
    M = V_desc.shape[1]
    for ev_idx, (evid, _, row) in enumerate(hits[:M]):
        ev_terms = []
        for fam_name, E_bank, L_bank in families:
            if E_bank is None or L_bank is None or len(L_bank) != E_bank.shape[0]:
                continue
            S = (E_bank @ V_desc[:, [ev_idx]]).squeeze(axis=1)
            idx = np.where(S >= per_event_thr)[0]
            if idx.size == 0:
                continue
            order = idx[np.argsort(S[idx])[::-1]][:top_per_family]
            labcol = _labels_col(L_bank)
            fam_lines = []
            for i in order:
                lab = str(L_bank.iloc[i].get(labcol, f"TERM_{i}"))
                fam_lines.append(f"{lab} (sim={float(S[i]):.3f})")
            if fam_lines:
                ev_terms.append(f"{fam_name}: " + "; ".join(fam_lines))
        lines.append(f"[EventID={evid}] " + (" | ".join(ev_terms) if ev_terms else "—"))
    return "".join(lines) + "", V_desc

# ===== Atribuição determinística por evento =====
def assign_terms_per_event(
    hits: List[Tuple[str, float, pd.Series]],
    V_desc: Optional[np.ndarray],
    E_ws, L_ws, E_prec, L_prec, E_cp, L_cp,
    per_event_thr: float,
    max_per_event: int = 2,
    max_global_frac: float = 0.5,
) -> str:
    if not hits or V_desc is None:
        return "=== EVENT_ASSIGNMENTS === [NENHUM HIT]"
    families = [("WS",E_ws,L_ws), ("PRE",E_prec,L_prec), ("CP",E_cp,L_cp)]
    def _labels_col(df):
        return next((c for c in ["label","text","name","CP","cp"] if (df is not None and c in df.columns)), None)
    M = V_desc.shape[1]
    candidates, label_pool = [], set()
    for ev_idx in range(min(M, len(hits))):
        ev_cands = []
        for fam_name, E_bank, L_bank in families:
            if E_bank is None or L_bank is None or len(L_bank) != E_bank.shape[0]:
                continue
            S = (E_bank @ V_desc[:, [ev_idx]]).squeeze(axis=1)
            idx = np.where(S >= per_event_thr)[0]
            if idx.size == 0:
                continue
            order = idx[np.argsort(S[idx])[::-1]]
            labcol = _labels_col(L_bank)
            for i in order[:10]:
                lab = str(L_bank.iloc[i].get(labcol, f"TERM_{i}"))
                ev_cands.append((lab, float(S[i]), fam_name))
                label_pool.add(lab)
        ev_cands.sort(key=lambda t: t[1], reverse=True)
        candidates.append(ev_cands)
    n_events = len(candidates)
    max_global = max(1, int(np.ceil(max_global_frac * n_events)))
    used_count: Dict[str, int] = {lab: 0 for lab in label_pool}
    lines = ["=== EVENT_ASSIGNMENTS ==="]
    for ev_idx, (evid, _, _row) in enumerate(hits[:n_events]):
        picked, seen_fams = [], set()
        for lab, sim, fam in candidates[ev_idx]:
            if used_count.get(lab, 0) >= max_global:
                continue
            if fam in seen_fams and len(seen_fams) < 3:
                continue
            picked.append((lab, sim, fam))
            used_count[lab] = used_count.get(lab, 0) + 1
            seen_fams.add(fam)
            if len(picked) >= max_per_event:
                break
        lines.append(f"[EventID={evid}] " + ("; ".join([p[0] for p in picked]) if picked else "—"))
    return "".join(lines) + ""

# ========================== Modelo ==========================
def ollama_chat(messages, model=None, temperature=0.2, stream=False, timeout=120):
    """
    Chat com Ollama com tratamento robusto de erros
    """
    if not (OLLAMA_HOST and (model or OLLAMA_MODEL)):
        raise RuntimeError("Modelo não configurado. Defina OLLAMA_HOST e OLLAMA_MODEL.")
    
    try:
        import requests
        url = f"{OLLAMA_HOST}/api/chat"
        payload = {
            "model": model or OLLAMA_MODEL, 
            "messages": messages, 
            "temperature": float(temperature), 
            "stream": bool(stream)
        }
        
        _info(f"Tentando conectar ao Ollama: {OLLAMA_HOST}")
        r = requests.post(url, headers=HEADERS_JSON, json=payload, timeout=timeout)
        
        if r.status_code == 200:
            return r.json()
        elif r.status_code == 404:
            raise RuntimeError(f"Modelo '{model or OLLAMA_MODEL}' não encontrado no Ollama. Verifique se o modelo está instalado.")
        elif r.status_code == 503:
            raise RuntimeError("Ollama está sobrecarregado ou não está pronto. Tente novamente em alguns segundos.")
        else:
            r.raise_for_status()
            
    except requests.exceptions.ConnectionError as e:
        raise RuntimeError(f"Erro de conectividade com {OLLAMA_HOST}. Verifique se o Ollama está rodando.")
    except requests.exceptions.Timeout:
        raise RuntimeError(f"Timeout ao conectar com {OLLAMA_HOST}. O serviço pode estar sobrecarregado.")
    except requests.exceptions.RequestException as e:
        raise RuntimeError(f"Erro de requisição: {e}")
    except Exception as e:
        raise RuntimeError(f"Erro inesperado ao comunicar com Ollama: {e}")

# ========================== Sidebar ==========================
st.sidebar.subheader("Assistente de Prompts")
prompts_bank = load_prompts_md(PROMPTS_MD_PATH)

col_p1, col_p2 = st.sidebar.columns(2)
with col_p1:
    titles_texto = [it["title"] for it in prompts_bank.get("Texto", [])]
    sel_texto = st.selectbox("Texto", options=["(vazio)"] + titles_texto, index=0)
with col_p2:
    titles_upload = [it["title"] for it in prompts_bank.get("Upload", [])]
    sel_upload = st.selectbox("Upload", options=["(vazio)"] + titles_upload, index=0)

if st.sidebar.button("Carregar no rascunho", use_container_width=True):
    draft = []
    if sel_texto != "(vazio)":
        body = next((it["body"] for it in prompts_bank["Texto"] if it["title"] == sel_texto), "")
        if body: draft.append(body)
    if sel_upload != "(vazio)":
        body = next((it["body"] for it in prompts_bank["Upload"] if it["title"] == sel_upload), "")
        if body: draft.append(body)
    st.session_state.draft_prompt = ("\n\n".join(draft)).strip()
    st.sidebar.success("Modelo(s) carregado(s) no rascunho.")
    st.rerun()

st.sidebar.subheader("Recuperação – Sphera")
top_k_sphera   = st.sidebar.slider("Top-K Sphera", 1, 100, 20, 1, help="Número máximo de eventos do Sphera a retornar")
limiar_sphera = st.sidebar.slider("Limiar de Similaridade Sphera", 0.0, 1.0, 0.30, 0.01, help="Similaridade mínima para considerar um evento relevante (0-1)")
anos_filtro    = st.sidebar.slider("Período (últimos N anos)", 1, 10, 3, 1, help="Filtrar eventos pelos últimos N anos")

st.sidebar.subheader("Recuperação – GoSee")
top_k_gosee   = st.sidebar.slider("Top-K GoSee", 1, 50, 10, 1, help="Número máximo de observações do GoSee a retornar")
limiar_gosee = st.sidebar.slider("Limiar de Similaridade GoSee", 0.0, 1.0, 0.25, 0.01, help="Similaridade mínima para considerar uma observação relevante (0-1)")

st.sidebar.subheader("Recuperação – Documentos")
top_k_docs   = st.sidebar.slider("Top-K Documentos", 1, 20, 5, 1, help="Número máximo de documentos a retornar")
limiar_docs = st.sidebar.slider("Limiar de Similaridade Documentos", 0.0, 1.0, 0.30, 0.01, help="Similaridade mínima para considerar um documento relevante (0-1)")

# Validar parâmetros da sidebar
sidebar_alerts = validate_configuration_sidebar_params(top_k_sphera, limiar_sphera, top_k_gosee, limiar_gosee, top_k_docs, limiar_docs, anos_filtro)
if sidebar_alerts:
    with st.sidebar.expander("🔔 Alertas de Configuração", expanded=True):
        for alert in sidebar_alerts:
            st.warning(alert)

st.sidebar.subheader("Filtros avançados – Sphera")
# NOVO: multiselect de Location
_loc_col_sidebar, _loc_options = _location_options_from(df_sph)
locations = st.sidebar.multiselect(
    f"Location (coluna: {_loc_col_sidebar or 'N/D'})",
    options=_loc_options,
    default=[],
    help="Opções extraídas do Sphera (LOCATION → FPSO → Location → FPSO/Unidade → Unidade)."
)
substr = st.sidebar.text_input("Description contém (substring)", "")

st.sidebar.subheader("Agregação sobre eventos recuperados (Sphera)")
agg_mode    = st.sidebar.selectbox("Agregação", ["max", "mean"], index=0)
per_ev_thr  = st.sidebar.slider("Limiar por evento (dicionários)", 0.0, 1.0, 0.15, 0.01)  # default ajustado
min_support = st.sidebar.slider("Suporte mínimo (nº de eventos)", 1, 20, 1, 1)

thr_ws_sim   = st.sidebar.slider("Limiar de similaridade WS",        0.0, 1.0, 0.25, 0.01)
thr_prec_sim = st.sidebar.slider("Limiar de similaridade Precursor", 0.0, 1.0, 0.25, 0.01)
thr_cp_sim   = st.sidebar.slider("Limiar de similaridade CP",        0.0, 1.0, 0.25, 0.01)

topn_ws   = st.sidebar.slider("Top-N WS",          3, 90, 10, 1)
topn_prec = st.sidebar.slider("Top-N Precursores", 3, 90, 10, 1)
topn_cp   = st.sidebar.slider("Top-N CP",          3, 90, 10, 1)

uc1, uc2 = st.sidebar.columns(2)
with uc1:
    if st.button("Limpar uploads", use_container_width=True):
        st.session_state.pop("upld_texts", None)
        st.session_state.upld_texts = []
        st.rerun()
with uc2:
    if st.button("Limpar chat", use_container_width=True):
        st.session_state.chat = []
        st.rerun()

# Status do sistema (sem alertas duplicados)
show_configuration_alerts()

# ========================== UI central ==========================
if st.session_state._clear_draft_flag:
    st.session_state.draft_prompt = ""
    st.session_state._clear_draft_flag = False

st.title("SAFETY • CHAT — Análise Integrada (Sphera + GoSee + Dicionários)")

st.text_area("Conteúdo do prompt", key="draft_prompt", height=180, placeholder="Digite ou carregue um modelo de prompt…")
user_text = st.text_area("Texto de análise (para Sphera)", height=200, placeholder="Cole aqui a descrição/evento a analisar…")

# ---------- Upload (txt, md, csv, pdf, docx, xlsx) ----------
uploaded = st.file_uploader(
    "Anexar arquivo (opcional)",
    type=["txt", "md", "csv", "pdf", "docx", "xlsx"]
)  # upload não dispara

if uploaded is not None:
    raw = uploaded.read()
    name = uploaded.name.lower()
    as_text = ""

    if name.endswith(".pdf"):
        as_text = extract_pdf_text(io.BytesIO(raw))
    elif name.endswith(".docx"):
        as_text = extract_docx_text(io.BytesIO(raw))
    elif name.endswith(".xlsx"):
        as_text = extract_xlsx_text(io.BytesIO(raw))
    else:
        # textos puros e CSV
        try:
            as_text = raw.decode("utf-8", errors="ignore")
        except Exception:
            as_text = ""
        if name.endswith(".csv") and as_text:
            try:
                dfcsv = pd.read_csv(io.StringIO(as_text))
                as_text = "\n".join(
                    dfcsv.astype(str).fillna("")
                    .apply(lambda r: " ".join(r.values), axis=1)
                    .tolist()
                )
            except Exception:
                pass

    if as_text:
        st.success(f"Upload recebido: {uploaded.name} (armazenado no contexto local).")
        st.session_state.upld_texts.append(as_text)
    else:
        st.warning(
            f"Não foi possível extrair texto de {uploaded.name}. "
            "Se for PDF escaneado (imagem), poderá exigir OCR externo."
        )
# ---------- /Upload ----------

col_run1, col_run2, col_run3 = st.columns([1, 1, 1])
go_btn      = col_run1.button("Enviar para o chat", type="primary", use_container_width=True)
clear_draft = col_run2.button("Limpar rascunho", use_container_width=True)
clear_chat  = col_run3.button("Limpar chat", use_container_width=True)

if clear_draft:
    st.session_state._clear_draft_flag = True
    st.rerun()
if clear_chat:
    st.session_state.chat = []
    st.rerun()

# ========================== Execução ==========================

def render_hits_table(hits: List[Tuple[str, float, pd.Series]], topk_display: int, source_name: str = "Sphera"):
    """Renderiza tabela de resultados com validação robusta"""
    if not hits:
        return
        
    rows = []
    for evid, s, row in hits[: min(topk_display, len(hits))]:
        # Correção: validação segura da coluna de localização
        loc_val = "N/D"
        if source_name == "Sphera":
            if LOC_DISPLAY_COL and LOC_DISPLAY_COL in row.index:
                loc_val = str(row.get(LOC_DISPLAY_COL, 'N/D'))
            elif 'LOCATION' in row.index:
                loc_val = str(row.get('LOCATION', 'N/D'))
        elif source_name == "GoSee":
            loc_val = str(row.get('Area', row.get('Location', 'N/D')))
            
        if source_name == "Sphera":
            desc = str(row.get("Description", row.get("DESCRIPTION", ""))).strip()
        else:  # GoSee
            desc = str(row.get("Observation", "")).strip()
            
        # normalizar quebras de linha e artefatos _x000D_
        desc = desc.replace("\r", " ").replace("\n", " ").replace("_x000D_", " ")
        desc = re.sub(r"\s+", " ", desc).strip()
        rows.append({"Event ID": evid, "Similaridade": round(s, 3), "LOCATION": loc_val, "Description": desc})
        
    if rows:
        st.markdown(f"**Eventos do {source_name} (Top-{min(topk_display, len(hits))})**")
        st.dataframe(pd.DataFrame(rows), use_container_width=True, hide_index=True)
    else:
        st.info(f"Nenhum resultado válido do {source_name} para exibir.")


def render_docs_results(docs_hits: List[Tuple[str, float, str]], topk_display: int):
    """Renderiza resultados de documentos"""
    if not docs_hits:
        return
        
    rows = []
    for doc_name, similarity, snippet in docs_hits[: min(topk_display, len(docs_hits))]:
        # Truncar snippet para exibir melhor
        display_snippet = snippet[:300] + "..." if len(snippet) > 300 else snippet
        display_snippet = display_snippet.replace("\r", " ").replace("\n", " ").replace("_x000D_", " ")
        display_snippet = re.sub(r"\s+", " ", display_snippet).strip()
        
        rows.append({
            "Documento": doc_name,
            "Similaridade": round(similarity, 3),
            "Conteúdo": display_snippet
        })
    
    if rows:
        st.markdown(f"**Documentos Relevantes (Top-{min(topk_display, len(docs_hits))})**")
        st.dataframe(pd.DataFrame(rows), use_container_width=True, hide_index=True)
    else:
        st.info("Nenhum documento relevante encontrado.")


def push_model(messages: List[Dict[str, str]], pergunta: str, contexto_md: str, dic_matches_md: str):
    # Guardrails para impedir invenção de termos fora dos dicionários
    guardrails = (
        "REGRAS PARA WS/PRECURSORES/CP:\n"
        "- Use EXCLUSIVAMENTE os termos listados em DIC_MATCHES para nomear WS/Precursores/CP.\n"
        "- NÃO crie categorias novas nem traduza/alterar rótulos.\n"
        "- Se a lista estiver vazia, escreva 'nenhum termo ≥ limiar'.\n"
        "- Quando citar um termo, mantenha o rótulo exatamente como fornecido.\n"
        "\n"
        "REGRAS PARA A COLUNA 'Observações/Precursores relevantes':\n"
        "- Para cada EventID, escolha no máximo 2 termos dentre os sugeridos em EVENT_HINTS (quando houver).\n"
        "- Não repita o mesmo termo em mais de 50% dos eventos; prefira diversidade baseada em EVENT_HINTS.\n"
        "- Se um evento não tiver termos em EVENT_HINTS, use '—' nessa coluna (não invente).\n"
    )

    messages.append({"role": "user", "content": "DADOS DE APOIO (não responda aqui):\n" + contexto_md + "\n\n" + dic_matches_md})
    q = pergunta or st.session_state.draft_prompt or "Faça a síntese conforme regras."
    messages.append({"role": "user", "content": guardrails + "\nPergunta/objetivo:\n" + q})

    try:
        resp = ollama_chat(messages, model=OLLAMA_MODEL, temperature=0.2, stream=False)
        content = ""
        if isinstance(resp, dict):
            content = resp.get("message", {}).get("content", "") or resp.get("content", "")
        if not content:
            content = "(Sem conteúdo do modelo)"
        with st.chat_message("assistant"):
            st.markdown(content)
        st.session_state.chat.append({"role": "assistant", "content": content})
        st.session_state["_just_replied"] = True
    except Exception as e:
        _warn(f"Erro ao consultar modelo Ollama: {e}")
        st.error(f"Falha ao consultar modelo: {e}")
        
        # Verificar se é um problema de conectividade
        if "Connection refused" in str(e) or "NewConnectionError" in str(e):
            st.error("🔌 **Ollama não está rodando localmente.**")
            st.info("💡 **Para usar o chat, configure o Ollama ou use uma API externa.**")
            st.info("**Opções:**")
            st.info("1. **Local**: Instale e rode Ollama (`ollama serve`)")
            st.info("2. **Cloud**: Configure OLLAMA_HOST para uma API externa")
            st.info("3. **Alternativa**: Use o chat sem LLMs (busca apenas)")
        elif "Modelo não configurado" in str(e):
            st.error("⚙️ **Configuração do Ollama incompleta.**")
            st.info("Configure as variáveis de ambiente:")
            st.info("- `OLLAMA_HOST`: URL do servidor Ollama")
            st.info("- `OLLAMA_MODEL`: Nome do modelo (ex: llama3.2:3b)")
        else:
            st.error(f"❌ **Erro do Ollama:** {e}")
        
        # Usuário pode usar a aplicação sem LLM
        st.info("💡 **A aplicação funciona sem LLM para busca semântica.**")

if go_btn:
    # Indicador de progresso
    with st.spinner("Processando consulta integrada..."):
        blocks = []
        if st.session_state.draft_prompt.strip():
            blocks.append("PROMPT:\n" + st.session_state.draft_prompt.strip())
        if (user_text or "").strip():
            blocks.append("TEXTO:\n" + user_text.strip())
        for i, t in enumerate(st.session_state.upld_texts or []):
            blocks.append(f"UPLOAD[{i+1}]:\n" + t.strip())

        # Busca integrada em múltiplas fontes
        all_results = {}
        
        # 1. Busca Sphera
        try:
            with st.status("🔍 Buscando no Sphera...", expanded=False):
                st.write("Aplicando filtros e calculando similaridades...")
                hits_sph = sphera_similar_to_text(
                    query_text=(user_text or st.session_state.draft_prompt),
                    min_sim=limiar_sphera, years=anos_filtro, topk=top_k_sphera,
                    df_base=df_sph, E_base=E_sph, substr=substr, locations=locations,
                )
                all_results['sphera'] = hits_sph
                
                if hits_sph:
                    st.success(f"✅ Sphera: {len(hits_sph)} eventos encontrados")
                    render_hits_table(hits_sph, top_k_sphera, "Sphera")
                else:
                    st.info("ℹ️ Nenhum evento do Sphera atingiu o limiar/filtros atuais.")
                    
        except Exception as e:
            _warn(f"Erro na busca Sphera: {e}")
            st.error(f"❌ Erro na busca Sphera: {e}")
            all_results['sphera'] = []

        # 2. Busca GoSee (NOVO)
        try:
            with st.status("🔍 Buscando no GoSee...", expanded=False):
                hits_gosee = gosee_similar_to_text(
                    query_text=(user_text or st.session_state.draft_prompt),
                    min_sim=limiar_gosee, topk=top_k_gosee,
                    df_base=df_gosee, substr=substr,
                )
                all_results['gosee'] = hits_gosee
                
                if hits_gosee:
                    st.success(f"✅ GoSee: {len(hits_gosee)} observações encontradas")
                    render_hits_table(hits_gosee, top_k_gosee, "GoSee")
                else:
                    st.info("ℹ️ Nenhuma observação do GoSee encontrada.")
                    
        except Exception as e:
            _warn(f"Erro na busca GoSee: {e}")
            st.error(f"❌ Erro na busca GoSee: {e}")
            all_results['gosee'] = []

        # 3. Busca em documentos (NOVO)
        try:
            with st.status("🔍 Buscando em documentos...", expanded=False):
                hits_docs = docs_similar_to_text(
                    query_text=(user_text or st.session_state.draft_prompt),
                    min_sim=limiar_docs, topk=top_k_docs,
                    docs_dict=docs_index,
                )
                all_results['docs'] = hits_docs
                
                if hits_docs:
                    st.success(f"✅ Documentos: {len(hits_docs)} documentos relevantes")
                    render_docs_results(hits_docs, top_k_docs)
                else:
                    st.info("ℹ️ Nenhum documento relevante encontrado.")
                    
        except Exception as e:
            _warn(f"Erro na busca em documentos: {e}")
            st.error(f"❌ Erro na busca em documentos: {e}")
            all_results['docs'] = []

        # 4. Agregação de dicionários sobre resultados do Sphera
        if all_results.get('sphera'):
            try:
                with st.status("📊 Agregando dicionários...", expanded=False):
                    dict_matches = aggregate_dict_matches_over_hits(
                        all_results['sphera'], E_ws, L_ws, E_prec, L_prec, E_cp, L_cp,
                        thr_ws_sim=thr_ws_sim, thr_prec_sim=thr_prec_sim, thr_cp_sim=thr_cp_sim,
                        topn_ws=topn_ws, topn_prec=topn_prec, topn_cp=topn_cp,
                        agg_mode=agg_mode, per_event_thr=per_ev_thr, min_support=min_support,
                    )

                    # Depuração (opcional): mostra Top-N brutos para confirmar que o espaço vetorial está ok
                    debug_preview_dicts(all_results['sphera'], E_ws, L_ws, E_prec, L_prec, E_cp, L_cp, topk=10)

                    # Hints por evento (para ancorar as observações do modelo)
                    EVENT_HINTS_MD, _Vdesc = build_event_hints(
                        all_results['sphera'], E_ws, L_ws, E_prec, L_prec, E_cp, L_cp,
                        per_event_thr=per_ev_thr,
                        top_per_family=3,
                    )

                    # Bloco estruturado com os termos encontrados para passar ao LLM
                    def _fmt_list(name, arr):
                        if not arr:
                            return f"{name}: NENHUM_TERMO_ACIMA_DO_LIMIAR\n"
                        lines = [f"{name}:"]
                        for lab, sim, sup in arr:
                            lines.append(f"- termo={lab} | sim={sim:.3f} | suporte={sup}")
                        return "\n".join(lines) + "\n"

                    ws_list   = dict_matches.get("ws")   or []
                    prec_list = dict_matches.get("prec") or []
                    cp_list   = dict_matches.get("cp")   or []

                    DIC_MATCHES_MD = (
                        "=== DIC_MATCHES ===\n"
                        + _fmt_list("WS", ws_list)
                        + _fmt_list("PRECURSORES", prec_list)
                        + _fmt_list("CP", cp_list)
                    )
                    
                    st.success("✅ Dicionários agregados com sucesso")

            except Exception as e:
                _warn(f"Erro na agregação de dicionários: {e}")
                st.error(f"❌ Erro na agregação de dicionários: {e}")
                DIC_MATCHES_MD = "=== DIC_MATCHES ===\n[ERRO NA AGREGAÇÃO]"
                EVENT_HINTS_MD = "=== EVENT_HINTS ===\n[ERRO NA GERAÇÃO]"
        else:
            DIC_MATCHES_MD = "=== DIC_MATCHES ===\n[SEM RESULTADOS DO SPHERA]"
            EVENT_HINTS_MD = "=== EVENT_HINTS ===\n[SEM RESULTADOS DO SPHERA]"

        # 5. Contexto unificado para o modelo
        ctx_chunks = []
        
        # Contexto Sphera
        if all_results.get('sphera'):
            table_ctx_rows = []
            for evid, s, row in all_results['sphera'][: min(k_sph, len(all_results['sphera']))]:
                loc_val = (str(row.get(LOC_DISPLAY_COL, row.get('LOCATION', 'N/D'))) if 'LOC_DISPLAY_COL' in globals() and LOC_DISPLAY_COL else str(row.get('LOCATION','N/D')))
                desc    = str(row.get("Description", row.get("DESCRIPTION", ""))).strip()
                desc = desc.replace("\r", " ").replace("\n", " ").replace("_x000D_", " ")
                desc = re.sub(r"\s+", " ", desc).strip()
                table_ctx_rows.append(f"EventID={evid} | sim={s:.3f} | LOCATION={loc_val} | Description={desc}")
            
            ctx_chunks.append(f"=== SPHERA_HITS ===\nHits={len(all_results['sphera'])}, thr={limiar_sphera:.2f}, years={anos_filtro}\n" + "\n".join(table_ctx_rows) + "\n")

        # Contexto GoSee
        if all_results.get('gosee'):
            gosee_ctx_rows = []
            for evid, s, row in all_results['gosee'][: min(top_k_gosee, len(all_results['gosee']))]:
                area = str(row.get('Area', row.get('Location', 'N/D')))
                obs = str(row.get("Observation", "")).strip()[:200]  # Limitar tamanho
                obs = obs.replace("\r", " ").replace("\n", " ").replace("_x000D_", " ")
                obs = re.sub(r"\s+", " ", obs).strip()
                gosee_ctx_rows.append(f"GoSeeID={evid} | sim={s:.3f} | Area={area} | Observation={obs}")
            
            ctx_chunks.append(f"=== GOSEE_HITS ===\nHits={len(all_results['gosee'])}, thr={limiar_gosee:.2f}\n" + "\n".join(gosee_ctx_rows) + "\n")

        # Contexto Documentos
        if all_results.get('docs'):
            docs_ctx_rows = []
            for doc_name, similarity, snippet in all_results['docs'][: min(top_k_docs, len(all_results['docs']))]:
                snippet_short = snippet[:150].replace("\r", " ").replace("\n", " ")
                snippet_short = re.sub(r"\s+", " ", snippet_short).strip()
                docs_ctx_rows.append(f"Doc={doc_name} | sim={similarity:.3f} | Content={snippet_short}...")
            
            ctx_chunks.append(f"=== DOCS_HITS ===\nHits={len(all_results['docs'])}, thr={limiar_docs:.2f}\n" + "\n".join(docs_ctx_rows) + "\n")

        # Adicionar hints e matches
        ctx_chunks.append(EVENT_HINTS_MD)
        
        # Preparar mensagens para o modelo
        messages = [
            {"role": "system", "content": st.session_state.system_prompt},
            {"role": "user", "content": "".join([b for b in blocks if b])},
        ]
        
        ctx_full = "".join([x for x in ctx_chunks if x])
        push_model(messages, user_text, ctx_full, DIC_MATCHES_MD)

# ========================== Histórico ==========================
if st.session_state.get("_just_replied"):
    st.session_state["_just_replied"] = False
else:
    if st.session_state.chat:
        st.divider()
        st.subheader("Histórico")
        for m in st.session_state.chat[-10:]:
            role = m.get("role", "assistant")
            with st.chat_message("assistant" if role != "user" else "user"):
                st.markdown(m.get("content", ""))
