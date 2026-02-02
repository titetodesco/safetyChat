
from __future__ import annotations
import io, pandas as pd

def extract_pdf_text(file) -> str:
    try:
        from pypdf import PdfReader
        # Tenta ler o PDF
        reader = PdfReader(file)
        text_parts = []
        
        # Se o PDF estiver criptografado, tenta sem senha
        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception:
                pass
        
        # Extrai texto de cada página
        for i, page in enumerate(reader.pages):
            try:
                text = page.extract_text()
                if text and text.strip():
                    text_parts.append(text.strip())
            except Exception as e:
                print(f"Erro na página {i+1}: {e}")
                continue
        
        if text_parts:
            return "\n\n".join(text_parts)
        else:
            print("PDF sem texto extraível (pode ser escaneado/imagem)")
            return ""
            
    except Exception as e:
        print(f"Erro ao abrir PDF: {type(e).__name__}: {str(e)}")
        return ""

def extract_docx_text(file) -> str:
    try:
        from docx import Document
        doc = Document(file)
        text_parts = []
        
        # Extrai texto dos parágrafos
        for para in doc.paragraphs:
            if para.text and para.text.strip():
                text_parts.append(para.text.strip())
        
        # Extrai texto das tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        text_parts.append(cell.text.strip())
        
        if text_parts:
            return "\n\n".join(text_parts)
        else:
            print("DOCX sem texto extraível")
            return ""
            
    except Exception as e:
        print(f"Erro ao abrir DOCX: {type(e).__name__}: {str(e)}")
        return ""

def extract_xlsx_text(file) -> str:
    try:
        dfs = pd.read_excel(file, sheet_name=None)
        parts = []
        for name, df in dfs.items():
            parts += df.astype(str).fillna("").apply(lambda r: " ".join(r.values), axis=1).tolist()
        
        if parts:
            return "\n".join(parts)
        else:
            print("XLSX sem dados extraíveis")
            return ""
            
    except Exception as e:
        print(f"Erro ao abrir XLSX: {type(e).__name__}: {str(e)}")
        return ""

def extract_csv_text(file) -> str:
    try:
        df = pd.read_csv(file)
        rows = df.astype(str).fillna("").apply(lambda r: " ".join(r.values), axis=1).tolist()
        
        if rows:
            return "\n".join(rows)
        else:
            print("CSV sem dados extraíveis")
            return ""
            
    except Exception as e:
        print(f"Erro ao abrir CSV: {type(e).__name__}: {str(e)}")
        return ""

def extract_any(uploaded_file) -> str:
    name = uploaded_file.name.lower()
    data = uploaded_file.read()
    bio = io.BytesIO(data)
    if name.endswith(".pdf"):
        return extract_pdf_text(bio)
    if name.endswith(".docx"):
        return extract_docx_text(bio)
    if name.endswith(".xlsx"):
        return extract_xlsx_text(bio)
    if name.endswith(".csv"):
        return extract_csv_text(bio)
    if name.endswith(".txt") or name.endswith(".md"):
        try:
            return data.decode("utf-8", errors="ignore")
        except Exception:
            return ""
    return ""
