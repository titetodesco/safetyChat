# -*- coding: utf-8 -*- v. 06/11/2025  20 hs
"""
app_chat_novo (1).py — Sphera + RAG + DIC (WS/Precursores/CP)

Patches desta versão:
- Top-K Sphera respeitado (sem cortes fixos em 10).
- WS/Precursores/CP calculados corretamente a partir dos hits do Sphera.
- Depuração opcional: Top-N brutos (ignora thresholds).
- Upload aceita txt, md, csv, pdf, docx, xlsx (sem OCR).
- Multiselect de Location na sidebar (LOCATION → FPSO → Location → FPSO/Unidade → Unidade; nunca AREA).
- Demais funcionalidades preservadas (prompts, contexto datasets, limpeza de estado, etc.).
"""

import os
import re
import io
from datetime import datetime, timedelta
from pathlib import Path
from typing import List, Dict, Tuple, Optional

import numpy as np
import pandas as pd
import streamlit as st

# Defina a sua senha secreta aqui
PASSWORD = "cdshell"  # Troque por uma senha forte

def check_password():
    """Exibe um campo de senha e retorna True se a senha estiver correta."""
    st.sidebar.header("🔒 Área protegida")
    password = st.sidebar.text_input("Digite a senha para acessar o app:", type="password")
    if password == PASSWORD:
        return True
    elif password:
        st.sidebar.error("Senha incorreta. Tente novamente.")
        return False
    else:
        return False

if not check_password():
    st.stop()  # Interrompe o app até digitar a senha correta

# ========================== Config ==========================
st.set_page_config(page_title="SAFETY • CHAT", page_icon="💬", layout="wide")

DATA_DIR = Path("data")
AN_DIR   = DATA_DIR / "analytics"
XLSX_DIR = DATA_DIR / "xlsx"
DATASETS_CONTEXT_PATH = DATA_DIR / "datasets_context.md"
PROMPTS_MD_PATH       = DATA_DIR / "prompts" / "prompts.md"

SPH_PQ_PATH  = AN_DIR / "sphera.parquet"
SPH_NPZ_PATH = AN_DIR / "sphera_embeddings.npz"

# Modelo (ajuste via secrets ou env)
OLLAMA_HOST    = st.secrets.get("OLLAMA_HOST", os.getenv("OLLAMA_HOST", ""))
OLLAMA_MODEL   = st.secrets.get("OLLAMA_MODEL", os.getenv("OLLAMA_MODEL", ""))
OLLAMA_API_KEY = st.secrets.get("OLLAMA_API_KEY", os.getenv("OLLAMA_API_KEY"))
HEADERS_JSON   = {"Authorization": f"Bearer {OLLAMA_API_KEY}", "Content-Type": "application/json"} if OLLAMA_API_KEY else {"Content-Type": "application/json"}

ST_MODEL_NAME = os.getenv("ST_MODEL_NAME", "sentence-transformers/all-MiniLM-L6-v2")

# ========================== Helpers básicos ==========================
def _fatal(msg: str):
    st.error(msg)
    st.stop()

try:
    from sentence_transformers import SentenceTransformer
except Exception as e:
    _fatal(f"❌ sentence-transformers indisponível: {e}")

@st.cache_resource(show_spinner=False)
def ensure_st_encoder():
    try:
        return SentenceTransformer(ST_MODEL_NAME)
    except Exception as e:
        _fatal(f"❌ Não foi possível carregar o encoder: {e}")

@st.cache_data(show_spinner=False)
def load_npz_embeddings(path: Path) -> Optional[np.ndarray]:
    if not path.exists():
        return None
    try:
        with np.load(str(path), allow_pickle=True) as z:
            for key in ("embeddings", "E", "X", "vectors", "vecs", "arr_0"):
                if key in z:
                    E = np.array(z[key]).astype(np.float32, copy=False)
                    E /= (np.linalg.norm(E, axis=1, keepdims=True) + 1e-9)
                    return E
            # fallback: maior matriz 2D
            best = None
            for k in z.files:
                arr = z[k]
                if isinstance(arr, np.ndarray) and arr.ndim == 2:
                    if best is None or arr.shape[0] > best.shape[0]:
                        best = arr
            if best is None:
                return None
            E = np.array(best).astype(np.float32, copy=False)
            E /= (np.linalg.norm(E, axis=1, keepdims=True) + 1e-9)
            return E
    except Exception:
        return None

@st.cache_data(show_spinner=False)
def load_prompts_md(md_path: Path) -> Dict[str, List[Dict[str, str]]]:
    if not md_path.exists():
        return {"Texto": [], "Upload": []}
    raw = md_path.read_text(encoding="utf-8")
    sections = re.split(r"(?m)^##\s+", raw)
    data = {"Texto": [], "Upload": []}
    for sec in sections:
        sec = sec.strip()
        if not sec:
            continue
        first, _, rest = sec.partition("\n")
        if first.strip() not in ("Texto", "Upload"):
            continue
        parts = re.split(r"(?m)^###\s+", rest)
        for p in parts:
            p = p.strip()
            if not p:
                continue
            title, _, body = p.partition("\n")
            data[first.strip()].append({"title": title.strip(), "body": body.strip()})
    return data

@st.cache_data(show_spinner=False)
def load_file_text(p: Path) -> str:
    try:
        return p.read_text(encoding="utf-8")
    except Exception as e:
        return f"[AVISO] Não consegui ler {p}: {e} (continuando sem este contexto)"

# ========================== Location utils ==========================
def get_sphera_location_col(df: pd.DataFrame) -> Optional[str]:
    """Retorna a melhor coluna de Location (nunca AREA).
    Prioridade: LOCATION → FPSO → Location → FPSO/Unidade → Unidade."""
    if df is None or df.empty:
        return None
    for c in ["LOCATION", "FPSO", "Location", "FPSO/Unidade", "Unidade"]:
        if c in df.columns and df[c].notna().any():
            return c
    return None

@st.cache_data(show_spinner=False)
def _location_options_from(df_full: pd.DataFrame) -> Tuple[Optional[str], List[str]]:
    col = get_sphera_location_col(df_full)
    if not col:
        return None, []
    s = df_full[col].astype(str).str.strip()
    s = s[(~s.isna()) & (s.str.len() > 0)]
    bad = {"nan", "none", "n/d", "nd"}
    s = s[~s.str.lower().isin(bad)]
    # de-duplicar preservando a primeira grafia
    seen = {}
    for v in s:
        k = v.lower()
        if k not in seen:
            seen[k] = v
    return col, sorted(seen.values())

# ========================== Carregamento de dados ==========================
if not SPH_PQ_PATH.exists():
    st.error(f"Parquet do Sphera não encontrado em {SPH_PQ_PATH}")

df_sph = pd.read_parquet(SPH_PQ_PATH) if SPH_PQ_PATH.exists() else pd.DataFrame()
E_sph  = load_npz_embeddings(SPH_NPZ_PATH)

# Dicionários
WS_NPZ,   WS_LBL   = AN_DIR / "ws_embeddings_pt.npz",   AN_DIR / "ws_embeddings_pt.parquet"
PREC_NPZ, PREC_LBL = AN_DIR / "prec_embeddings_pt.npz", AN_DIR / "prec_embeddings_pt.parquet"
CP_NPZ,   CP_LBL   = AN_DIR / "cp_embeddings.npz",      AN_DIR / "cp_labels.parquet"

E_ws   = load_npz_embeddings(WS_NPZ) if WS_NPZ.exists() else None
L_ws   = (pd.read_parquet(WS_LBL) if WS_LBL.exists() else None)
E_prec = load_npz_embeddings(PREC_NPZ) if PREC_NPZ.exists() else None
L_prec = (pd.read_parquet(PREC_LBL) if PREC_LBL.exists() else None)
E_cp   = load_npz_embeddings(CP_NPZ) if CP_NPZ.exists() else None
L_cp   = (pd.read_parquet(CP_LBL) if CP_LBL.exists() else None)

# ========================== Estado ==========================
if "system_prompt" not in st.session_state:
    pre = (
        "Você é o ESO-CHAT para segurança operacional (óleo e gás). "
        "Responda em PT-BR, cite IDs/sim quando usar buscas locais, e não invente dados fora dos contextos fornecidos.\n\n"
    )
    sys_ctx = (load_file_text(DATASETS_CONTEXT_PATH) if DATASETS_CONTEXT_PATH.exists() else "")
    st.session_state.system_prompt = pre + ("=== DATASETS_CONTEXT ===\n" + sys_ctx if sys_ctx else "")
if "chat" not in st.session_state:
    st.session_state.chat = []
if "draft_prompt" not in st.session_state:
    st.session_state.draft_prompt = ""
if "_clear_draft_flag" not in st.session_state:
    st.session_state._clear_draft_flag = False
if "st_encoder" not in st.session_state:
    st.session_state.st_encoder = ensure_st_encoder()
if "upld_texts" not in st.session_state:
    st.session_state.upld_texts = []

# ========================== Encode ==========================
@st.cache_data(show_spinner=False)
def encode_texts(texts: List[str], batch_size: int = 64) -> np.ndarray:
    M = st.session_state.st_encoder.encode(
        texts, batch_size=batch_size, show_progress_bar=False,
        convert_to_numpy=True, normalize_embeddings=True
    ).astype(np.float32)
    return M

@st.cache_data(show_spinner=False)
def encode_query(q: str) -> np.ndarray:
    v = st.session_state.st_encoder.encode([q], convert_to_numpy=True, normalize_embeddings=True)[0].astype(np.float32)
    v /= (np.linalg.norm(v) + 1e-9)
    return v

# ========================== Filtros / Similaridade ==========================
@st.cache_data(show_spinner=False)
def filter_sphera(df: pd.DataFrame, locations: List[str], substr: str, years: int) -> pd.DataFrame:
    if df is None or df.empty:
        return pd.DataFrame()
    out = df.copy()

    # Janela temporal
    if "EVENT_DATE" in out.columns:
        out["EVENT_DATE"] = pd.to_datetime(out["EVENT_DATE"], errors="coerce")
        cutoff = pd.Timestamp(datetime.utcnow() - timedelta(days=365 * years))
        out = out[out["EVENT_DATE"] >= cutoff]

    # Filtro por Location
    loc_col = get_sphera_location_col(out)
    if loc_col and locations:
        sel = set([str(x).strip() for x in locations if str(x).strip()])
        out = out[out[loc_col].astype(str).isin(sel)]

    # Description contém (case-insensitive)
    desc_col = "Description" if "Description" in out.columns else ("DESCRIPTION" if "DESCRIPTION" in out.columns else None)
    if desc_col and substr:
        pat = re.escape(substr)
        out = out[out[desc_col].astype(str).str.contains(pat, case=False, na=False, regex=True)]

    return out

@st.cache_data(show_spinner=False)
def sphera_similar_to_text(query_text: str, min_sim: float, years: int, topk: int,
                           df_base: pd.DataFrame, E_base: Optional[np.ndarray],
                           substr: str, locations: List[str]) -> List[Tuple[str, float, pd.Series]]:
    if not query_text or df_base is None or df_base.empty or E_base is None or E_base.size == 0:
        return []
    base = filter_sphera(df_base, locations, substr, years)
    if base.empty:
        return []
    # tentar alinhar índices do parquet com os embeddings
    try:
        idx_map = base.index.to_numpy()
        if np.issubdtype(idx_map.dtype, np.integer):
            E_view = E_base[idx_map, :]
        else:
            E_view = E_base
            base = df_base
    except Exception:
        E_view = E_base
        base = df_base
    qv = encode_query(query_text)
    sims = (E_view @ qv).astype(float)
    ord_idx = np.argsort(-sims)
    id_col = next((c for c in ["Event ID", "EVENT_NUMBER", "EVENTID"] if c in base.columns), None)
    out = []
    kept = 0
    for i in ord_idx:
        s = float(sims[i])
        if s < min_sim:
            continue
        row = base.iloc[int(i)]
        evid = row.get(id_col, f"row{i}") if id_col else f"row{i}"
        out.append((str(evid), s, row))
        kept += 1
        if kept >= topk:
            break
    return out

# ========================== Agregação dicionários ==========================
@st.cache_data(show_spinner=False)
def aggregate_dict_matches_over_hits(
    hits: List[Tuple[str, float, pd.Series]],
    E_ws, L_ws, E_prec, L_prec, E_cp, L_cp,
    thr_ws_sim: float, thr_prec_sim: float, thr_cp_sim: float,
    topn_ws: int, topn_prec: int, topn_cp: int,
    agg_mode: str = "max",
    per_event_thr: float = 0.15,
    min_support: int = 1,
) -> Dict[str, List[Tuple[str, float, int]]]:
    if not hits:
        return {"ws": [], "prec": [], "cp": []}

    # Descrições dos HITS (não do DF completo)
    descs = [str(r.get("Description", r.get("DESCRIPTION", ""))).strip() for _, _, r in hits]
    descs = [d for d in descs if d]
    if not descs:
        return {"ws": [], "prec": [], "cp": []}

    # Embeddings das descrições (normalizados pelo encoder)
    V_desc = encode_texts(descs, batch_size=32).T  # (D x M)

    def _labels_col(df: pd.DataFrame) -> Optional[str]:
        return next((c for c in ["label", "text", "name", "CP", "cp"] if c in df.columns), None)

    def _score(E_bank, labels_df, thr_sim, topn_target):
        if E_bank is None or labels_df is None or len(labels_df) != E_bank.shape[0]:
            return []
        S = (E_bank @ V_desc)  # (N_terms x M_events) — cos sim
        support = (S >= per_event_thr).sum(axis=1)
        sims = S.mean(axis=1) if agg_mode == "mean" else S.max(axis=1)

        mask = (support >= min_support) & (sims >= thr_sim)
        idx = np.where(mask)[0]
        if idx.size == 0:
            return []

        order = idx[np.argsort(sims[idx])[::-1]]
        labcol = _labels_col(labels_df)
        out = []
        for i in order[:topn_target]:
            label = str(labels_df.iloc[i].get(labcol, f"TERM_{i}"))
            out.append((label, float(sims[i]), int(support[i])))
        return out

    return {
        "ws":   _score(E_ws,   L_ws,   thr_ws_sim,   topn_ws),
        "prec": _score(E_prec, L_prec, thr_prec_sim, topn_prec),
        "cp":   _score(E_cp,   L_cp,   thr_cp_sim,   topn_cp),
    }

# ===== Depuração (opcional): Top-N "brutos" (ignora thresholds) =====
def _topk_raw_for_bank(E_bank, labels_df, V_desc_T, topk=10):
    if E_bank is None or labels_df is None or len(labels_df) != (E_bank.shape[0] if hasattr(E_bank, "shape") else 0) or V_desc_T is None:
        return pd.DataFrame()
    S = (E_bank @ V_desc_T)
    sims = S.max(axis=1)
    order = np.argsort(sims)[::-1][:topk]
    labcol = next((c for c in ["label", "text", "name", "CP", "cp"] if c in labels_df.columns), "label")
    rows = [{"Termo": str(labels_df.iloc[i].get(labcol, f"TERM_{i}")), "Similaridade(max)": float(sims[i])} for i in order]
    return pd.DataFrame(rows)

def debug_preview_dicts(hits, E_ws, L_ws, E_prec, L_prec, E_cp, L_cp, topk=10):
    if not hits:
        return
    descs = [str(r.get("Description", r.get("DESCRIPTION",""))).strip() for _,_,r in hits]
    descs = [d for d in descs if d]
    if not descs:
        return
    V_desc = encode_texts(descs, batch_size=32).T  # (D x M)
    with st.expander("🔎 Depuração — Top-N brutos (ignora thresholds)", expanded=False):
        if E_ws is not None and L_ws is not None:
            st.markdown("**WS**")
            st.dataframe(_topk_raw_for_bank(E_ws, L_ws, V_desc, topk), use_container_width=True, hide_index=True)
        if E_prec is not None and L_prec is not None:
            st.markdown("**Precursores**")
            st.dataframe(_topk_raw_for_bank(E_prec, L_prec, V_desc, topk), use_container_width=True, hide_index=True)
        if E_cp is not None and L_cp is not None:
            st.markdown("**CP**")
            st.dataframe(_topk_raw_for_bank(E_cp, L_cp, V_desc, topk), use_container_width=True, hide_index=True)

# ========================== Modelo ==========================
def ollama_chat(messages, model=None, temperature=0.2, stream=False, timeout=120):
    if not (OLLAMA_HOST and (model or OLLAMA_MODEL)):
        raise RuntimeError("Modelo não configurado. Defina OLLAMA_HOST e OLLAMA_MODEL.")
    import requests
    r = requests.post(f"{OLLAMA_HOST}/api/chat", headers=HEADERS_JSON, json={
        "model": model or OLLAMA_MODEL, "messages": messages, "temperature": float(temperature), "stream": bool(stream)
    }, timeout=timeout)
    r.raise_for_status()
    return r.json()

# ========================== Sidebar ==========================
st.sidebar.subheader("Assistente de Prompts")
prompts_bank = load_prompts_md(PROMPTS_MD_PATH)
c1, c2 = st.sidebar.columns(2)
with c1:
    titles_texto = [it["title"] for it in prompts_bank.get("Texto", [])]
    sel_texto = st.selectbox("Texto", options=["(vazio)"] + titles_texto, index=0)
with c2:
    titles_upload = [it["title"] for it in prompts_bank.get("Upload", [])]
    sel_upload = st.selectbox("Upload", options=["(vazio)"] + titles_upload, index=0)

if st.sidebar.button("Carregar no rascunho", use_container_width=True):
    draft = []
    if sel_texto != "(vazio)":
        body = next((it["body"] for it in prompts_bank["Texto"] if it["title"] == sel_texto), "")
        if body: draft.append(body)
    if sel_upload != "(vazio)":
        body = next((it["body"] for it in prompts_bank["Upload"] if it["title"] == sel_upload), "")
        if body: draft.append(body)
    st.session_state.draft_prompt = ("\n\n".join([d for d in draft if d])).strip()
    st.sidebar.success("Modelo(s) carregado(s) no rascunho.")
    st.rerun()

st.sidebar.header("Recuperação – Sphera")
k_sph   = st.sidebar.slider("Top-K Sphera", 1, 100, 20, 1)
thr_sph = st.sidebar.slider("Limiar Sphera (cos)", 0.0, 1.0, 0.30, 0.01)
years   = st.sidebar.slider("Últimos N anos", 1, 10, 3, 1)

st.sidebar.subheader("Filtros avançados – Sphera")
_loc_col_sidebar, _loc_options = _location_options_from(df_sph)
locations = st.sidebar.multiselect(
    f"Location (coluna: {_loc_col_sidebar or 'N/D'})",
    options=_loc_options,
    default=[],
    help="Lista derivada do dataframe (LOCATION → FPSO → Location → FPSO/Unidade → Unidade)."
)
substr = st.sidebar.text_input("Description contém (substring)", "")

st.sidebar.subheader("Agregação sobre eventos recuperados (Sphera)")
agg_mode    = st.sidebar.selectbox("Agregação", ["max", "mean"], index=0)
per_ev_thr  = st.sidebar.slider("Limiar por evento (dicionários)", 0.0, 1.0, 0.15, 0.01)
min_support = st.sidebar.slider("Suporte mínimo (nº de eventos)", 1, 20, 1, 1)

thr_ws_sim   = st.sidebar.slider("Limiar de similaridade WS",        0.0, 1.0, 0.25, 0.01)
thr_prec_sim = st.sidebar.slider("Limiar de similaridade Precursor", 0.0, 1.0, 0.25, 0.01)
thr_cp_sim   = st.sidebar.slider("Limiar de similaridade CP",        0.0, 1.0, 0.25, 0.01)

topn_ws   = st.sidebar.slider("Top-N WS",          3, 90, 10, 1)
topn_prec = st.sidebar.slider("Top-N Precursores", 3, 90, 10, 1)
topn_cp   = st.sidebar.slider("Top-N CP",          3, 90, 10, 1)

uc1, uc2 = st.sidebar.columns(2)
with uc1:
    if st.button("Limpar uploads", use_container_width=True):
        st.session_state.pop("upld_texts", None)
        st.session_state.upld_texts = []
        st.rerun()
with uc2:
    if st.button("Limpar chat", use_container_width=True):
        st.session_state.chat = []
        st.rerun()

# ========================== UI central ==========================
if st.session_state._clear_draft_flag:
    st.session_state.draft_prompt = ""
    st.session_state._clear_draft_flag = False

st.title("SAFETY • CHAT")

st.text_area("Conteúdo do prompt", key="draft_prompt", height=180, placeholder="Digite ou carregue um modelo de prompt…")
user_text = st.text_area("Texto de análise (para Sphera)", height=200, placeholder="Cole aqui a descrição/evento a analisar…")

# ---------- Upload (txt, md, csv, pdf, docx, xlsx) ----------
uploaded = st.file_uploader(
    "Anexar arquivo (opcional)",
    type=["txt", "md", "csv", "pdf", "docx", "xlsx"]
)  # upload não dispara

def extract_pdf_text(file_like: io.BytesIO) -> str:
    """Extrai texto de PDF. Tenta PyPDF2 -> PyMuPDF -> pdfminer.six; sem OCR."""
    # 1) PyPDF2
    try:
        import PyPDF2
        file_like.seek(0)
        reader = PyPDF2.PdfReader(file_like)
        parts = [(page.extract_text() or "") for page in reader.pages]
        return "\n".join(parts).strip()
    except Exception:
        pass
    # 2) PyMuPDF
    try:
        import fitz  # PyMuPDF
        file_like.seek(0)
        doc = fitz.open(stream=file_like.read(), filetype="pdf")
        parts = [page.get_text() for page in doc]
        return "\n".join(parts).strip()
    except Exception:
        pass
    # 3) pdfminer.six
    try:
        from pdfminer.high_level import extract_text
        file_like.seek(0)
        return (extract_text(file_like) or "").strip()
    except Exception:
        pass
    return ""

def extract_docx_text(file_like: io.BytesIO) -> str:
    """Extrai texto de um .docx (python-docx)."""
    try:
        from docx import Document
        file_like.seek(0)
        doc = Document(file_like)
        parts = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" ".join(cell.text for cell in row.cells if cell.text))
        return "\n".join(parts).strip()
    except Exception:
        return ""

def extract_xlsx_text(file_like: io.BytesIO) -> str:
    """Extrai texto de um .xlsx (pandas + openpyxl)."""
    try:
        file_like.seek(0)
        sheets = pd.read_excel(file_like, sheet_name=None, engine="openpyxl")
        lines = []
        for name, df in sheets.items():
            if df is None or df.empty:
                continue
            df = df.astype(str).fillna("")
            lines.append(f"=== SHEET: {name} ===")
            lines.extend(df.apply(lambda r: " ".join(r.values), axis=1).tolist())
        return "\n".join(lines).strip()
    except Exception:
        return ""

if uploaded is not None:
    raw = uploaded.read()
    name = uploaded.name.lower()
    as_text = ""

    if name.endswith(".pdf"):
        as_text = extract_pdf_text(io.BytesIO(raw))
    elif name.endswith(".docx"):
        as_text = extract_docx_text(io.BytesIO(raw))
    elif name.endswith(".xlsx"):
        as_text = extract_xlsx_text(io.BytesIO(raw))
    else:
        # textos puros e CSV
        try:
            as_text = raw.decode("utf-8", errors="ignore")
        except Exception:
            as_text = ""
        if name.endswith(".csv") and as_text:
            try:
                dfcsv = pd.read_csv(io.StringIO(as_text))
                as_text = "\n".join(
                    dfcsv.astype(str).fillna("")
                    .apply(lambda r: " ".join(r.values), axis=1)
                    .tolist()
                )
            except Exception:
                pass

    if as_text:
        st.success(f"Upload recebido: {uploaded.name} (armazenado no contexto local).")
        st.session_state.upld_texts.append(as_text)
    else:
        st.warning(
            f"Não foi possível extrair texto de {uploaded.name}. "
            "Se for PDF escaneado (imagem), poderá exigir OCR externo."
        )
# ---------- /Upload ----------

col_run1, col_run2, col_run3 = st.columns([1, 1, 1])
go_btn      = col_run1.button("Enviar para o chat", type="primary", use_container_width=True)
clear_draft = col_run2.button("Limpar rascunho", use_container_width=True)
clear_chat  = col_run3.button("Limpar chat", use_container_width=True)

if clear_draft:
    st.session_state._clear_draft_flag = True
    st.rerun()
if clear_chat:
    st.session_state.chat = []
    st.rerun()

# ========================== Execução ==========================
LOC_DISPLAY_COL = get_sphera_location_col(df_sph)

def render_hits_table(hits: List[Tuple[str, float, pd.Series]], topk_display: int):
    if not hits:
        return
    rows = []
    for evid, s, row in hits[: min(topk_display, len(hits))]:
        loc_val = (str(row.get(LOC_DISPLAY_COL, row.get("LOCATION", "N/D"))) if LOC_DISPLAY_COL else str(row.get("LOCATION","N/D")))
        desc    = str(row.get("Description", row.get("DESCRIPTION", ""))).strip()
        # normalização leve de quebras/artefatos
        desc = desc.replace("\r", " ").replace("\n", " ").replace("_x000D_", " ")
        desc = re.sub(r"\s+", " ", desc).strip()
        rows.append({"Event ID": evid, "Similaridade": round(s, 3), "LOCATION": loc_val, "Description": desc})
    st.markdown(f"**Eventos do Sphera (Top-{min(topk_display, len(hits))})**")
    st.dataframe(pd.DataFrame(rows), use_container_width=True, hide_index=True)

def push_model(messages: List[Dict[str, str]], pergunta: str, contexto_md: str, dic_matches_md: str):
    # Guardrails para o LLM não inventar categorias
    guardrails = (
        "REGRAS PARA WS/PRECURSORES/CP:\n"
        "- Use EXCLUSIVAMENTE os termos listados em DIC_MATCHES para nomear WS/Precursores/CP.\n"
        "- NÃO crie categorias novas nem traduza/alterar rótulos.\n"
        "- Se a lista estiver vazia, escreva 'nenhum termo ≥ limiar'.\n"
    )
    messages.append({"role": "user", "content": "DADOS DE APOIO (não responda aqui):\n" + contexto_md + "\n\n" + dic_matches_md})
    q = pergunta or st.session_state.draft_prompt or "Faça a síntese conforme regras."
    messages.append({"role": "user", "content": guardrails + "\nPergunta/objetivo:\n" + q})
    try:
        resp = ollama_chat(messages, model=OLLAMA_MODEL, temperature=0.2, stream=False)
        content = ""
        if isinstance(resp, dict):
            content = resp.get("message", {}).get("content", "") or resp.get("content", "")
        if not content:
            content = "(Sem conteúdo do modelo)"
        with st.chat_message("assistant"):
            st.markdown(content)
        st.session_state.chat.append({"role": "assistant", "content": content})
        st.session_state["_just_replied"] = True
    except Exception as e:
        st.error(f"Falha ao consultar modelo: {e}")

if go_btn:
    blocks = []
    if st.session_state.draft_prompt.strip():
        blocks.append("PROMPT:\n" + st.session_state.draft_prompt.strip())
    if (user_text or "").strip():
        blocks.append("TEXTO:\n" + user_text.strip())
    for i, t in enumerate(st.session_state.upld_texts or []):
        blocks.append(f"UPLOAD[{i+1}]:\n" + t.strip())

    # ===== Recuperação Sphera (Top-K correto) =====
    hits = sphera_similar_to_text(
        query_text=(user_text or st.session_state.draft_prompt),
        min_sim=thr_sph, years=years, topk=k_sph,
        df_base=df_sph, E_base=E_sph, substr=substr, locations=locations,
    )

    if hits:
        render_hits_table(hits, k_sph)
    else:
        st.info("Nenhum evento do Sphera atingiu o limiar/filtros atuais.")

    # ===== Agregação de dicionários nos hits =====
    dict_matches = aggregate_dict_matches_over_hits(
        hits, E_ws, L_ws, E_prec, L_prec, E_cp, L_cp,
        thr_ws_sim=thr_ws_sim, thr_prec_sim=thr_prec_sim, thr_cp_sim=thr_cp_sim,
        topn_ws=topn_ws, topn_prec=topn_prec, topn_cp=topn_cp,
        agg_mode=agg_mode, per_event_thr=per_ev_thr, min_support=min_support,
    )

    # Depuração opcional — ver os Top-N “brutos”
    debug_preview_dicts(hits, E_ws, L_ws, E_prec, L_prec, E_cp, L_cp, topk=10)

    # Bloco estruturado com termos encontrados (para o LLM não inventar)
    def _fmt_list(name, arr):
        if not arr:
            return f"{name}: NENHUM_TERMO_ACIMA_DO_LIMIAR\n"
        lines = [f"{name}:"]
        for lab, sim, sup in arr:
            lines.append(f"- termo={lab} | sim={sim:.3f} | suporte={sup}")
        return "\n".join(lines) + "\n"

    ws_list   = dict_matches.get("ws")   or []
    prec_list = dict_matches.get("prec") or []
    cp_list   = dict_matches.get("cp")   or []

    DIC_MATCHES_MD = (
        "=== DIC_MATCHES ===\n"
        + _fmt_list("WS", ws_list)
        + _fmt_list("PRECURSORES", prec_list)
        + _fmt_list("CP", cp_list)
    )

    # Contexto Sphera textual (Top-K)
    table_ctx_rows = []
    for evid, s, row in hits[: min(k_sph, len(hits))]:
        loc_val = (str(row.get(LOC_DISPLAY_COL, row.get("LOCATION", "N/D"))) if LOC_DISPLAY_COL else str(row.get("LOCATION","N/D")))
        desc    = str(row.get("Description", row.get("DESCRIPTION", ""))).strip()
        desc = desc.replace("\r", " ").replace("\n", " ").replace("_x000D_", " ")
        desc = re.sub(r"\s+", " ", desc).strip()
        table_ctx_rows.append(f"EventID={evid} | sim={s:.3f} | LOCATION={loc_val} | Description={desc}")

    ctx_chunks = [
        f"Sphera_hits={len(hits)}, thr_sph={thr_sph:.2f}, years={years}",
        "\n".join(table_ctx_rows),
    ]

    messages = [
        {"role": "system", "content": st.session_state.system_prompt},
        {"role": "user",   "content": "\n\n".join([b for b in blocks if b])},
    ]
    ctx_full = "\n\n".join([x for x in ctx_chunks if x])
    push_model(messages, user_text, ctx_full, DIC_MATCHES_MD)

# ========================== Histórico ==========================
if st.session_state.get("_just_replied"):
    st.session_state["_just_replied"] = False
else:
    if st.session_state.chat:
        st.divider()
        st.subheader("Histórico")
        for m in st.session_state.chat[-10:]:
            role = m.get("role", "assistant")
            with st.chat_message("assistant" if role != "user" else "user"):
                st.markdown(m.get("content", ""))
